from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path(r"sistema\FICHA-PRODUTO-MAIS-TECH\REFORCO_LINGUAGENS\APOSTILA_Reforco_Linguagens.docx")

BLUE = "1F4E79"
LIGHT_BLUE = "E8EEF5"
DARK = "203040"
GOLD = "C69214"
LIGHT_GOLD = "FFF4D6"
GRAY = "5B6573"
WHITE = "FFFFFF"


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_widths(table, widths):
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)
            set_cell_margins(row.cells[idx])
            row.cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def font(run, name="Calibri", size=11, color=DARK, bold=False, italic=False):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic
    return run


def add_para(doc, text="", *, bold=False, italic=False, color=DARK, size=11,
             align=None, before=0, after=6, line=1.25, keep=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    p.paragraph_format.keep_with_next = keep
    if align is not None:
        p.alignment = align
    font(p.add_run(text), size=size, color=color, bold=bold, italic=italic)
    return p


def add_callout(doc, label, text, fill=LIGHT_BLUE, accent=BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Inches(6.5)
    shade(cell, fill)
    set_cell_margins(cell, 130, 170, 130, 170)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.2
    font(p.add_run(f"{label}: "), size=10.5, color=accent, bold=True)
    font(p.add_run(text), size=10.5, color=DARK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.2
        font(p.add_run(item), size=11)


def add_activity(doc, title, objective, steps, product, time="30 a 50 minutos"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    font(p.add_run(title), size=12, color=BLUE, bold=True)
    add_callout(doc, "Objetivo", objective, LIGHT_GOLD, GOLD)
    add_para(doc, f"Tempo sugerido: {time}", italic=True, color=GRAY, size=9.5, after=4)
    for step in steps:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        font(p.add_run(step), size=10.5)
    add_para(doc, f"Entrega: {product}", bold=True, color=BLUE, size=10.5, after=8)


def add_module_intro(doc, number, title, hours, goals):
    doc.add_page_break()
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(48)
    p.paragraph_format.space_after = Pt(8)
    font(p.add_run(f"MÓDULO {number}"), size=12, color=GOLD, bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    font(p.add_run(title), size=25, color=BLUE, bold=True)
    add_para(doc, f"Carga horária: {hours}", bold=True, color=GRAY, size=11, after=16)
    add_callout(doc, "Neste módulo", goals)


def add_lesson(doc, title, explanation, example, practice):
    doc.add_heading(title, level=2)
    add_para(doc, explanation)
    add_callout(doc, "Exemplo", example)
    add_activity(doc, f"Atividade - {title}", practice[0], practice[1], practice[2], practice[3])
    add_para(doc, "Síntese: explique com suas palavras o que aprendeu e registre uma dúvida que ainda possui.",
             italic=True, color=GRAY, size=10, after=10)


doc = Document()
sec = doc.sections[0]
sec.page_width = Inches(8.5)
sec.page_height = Inches(11)
sec.top_margin = Inches(1)
sec.bottom_margin = Inches(1)
sec.left_margin = Inches(1)
sec.right_margin = Inches(1)
sec.header_distance = Inches(0.492)
sec.footer_distance = Inches(0.492)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(11)
normal.font.color.rgb = RGBColor.from_string(DARK)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.25

for name, size, color, before, after in (
    ("Heading 1", 16, BLUE, 18, 10),
    ("Heading 2", 13, BLUE, 14, 7),
    ("Heading 3", 12, "1F4D78", 10, 5),
):
    st = styles[name]
    st.font.name = "Calibri"
    st._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    st._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = RGBColor.from_string(color)
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.keep_with_next = True

header = sec.header
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
font(hp.add_run("RIO DO SUL MAIS TECH  |  REFORÇO DE LINGUAGENS"), size=8.5, color=GRAY, bold=True)

footer = sec.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
font(fp.add_run("SENAI | Prefeitura Municipal de Rio do Sul  •  "), size=8.5, color=GRAY)
fld = OxmlElement("w:fldSimple")
fld.set(qn("w:instr"), "PAGE")
fp._p.append(fld)

# Capa editorial
add_para(doc, "RIO DO SUL MAIS TECH", bold=True, color=GOLD, size=11,
         align=WD_ALIGN_PARAGRAPH.CENTER, before=90, after=20)
add_para(doc, "APOSTILA DO ALUNO", bold=True, color=GRAY, size=10,
         align=WD_ALIGN_PARAGRAPH.CENTER, after=10)
add_para(doc, "Reforço de Linguagens", bold=True, color=BLUE, size=30,
         align=WD_ALIGN_PARAGRAPH.CENTER, after=12)
add_para(doc, "Leitura, escrita, gramática, comunicação oral e autonomia para aprender",
         color=GRAY, size=14, align=WD_ALIGN_PARAGRAPH.CENTER, after=38)
add_callout(doc, "Público", "Estudantes do 8º e 9º ano do Ensino Fundamental (12 a 15 anos).")
add_para(doc, "Carga horária total: 63 horas", bold=True, color=BLUE, size=12,
         align=WD_ALIGN_PARAGRAPH.CENTER, before=48, after=8)
add_para(doc, "SENAI | Prefeitura Municipal de Rio do Sul", color=GRAY, size=10,
         align=WD_ALIGN_PARAGRAPH.CENTER, after=0)

doc.add_page_break()
doc.add_heading("Apresentação", level=1)
add_para(doc, "Esta apostila foi criada para fortalecer sua capacidade de compreender textos, organizar ideias, escrever com clareza e comunicar-se com segurança. A língua está presente em todas as áreas: nas mensagens, nos estudos, nas entrevistas, nos projetos e nas relações do dia a dia.")
add_para(doc, "Você encontrará explicações diretas, exemplos próximos da sua realidade e atividades individuais e colaborativas. Errar faz parte do processo: cada tentativa ajuda a perceber o que já foi aprendido e qual será o próximo passo.")
add_callout(doc, "Como estudar", "Leia com atenção, marque palavras importantes, faça anotações, compare respostas e sempre revise antes de entregar.", LIGHT_GOLD, GOLD)

doc.add_heading("Objetivo da unidade curricular", level=1)
add_para(doc, "Desenvolver e aprimorar as competências de leitura, escrita, interpretação e comunicação oral, fortalecendo a base linguística para a vida acadêmica, profissional e social.")
doc.add_heading("Competências que você desenvolverá", level=2)
add_bullets(doc, [
    "Organizar informações de maneira clara, coerente e adequada ao objetivo.",
    "Utilizar estratégias de leitura e escrita com maior autonomia.",
    "Analisar textos e situações comunicativas com pensamento crítico.",
    "Trabalhar em grupo, escutar opiniões e colaborar com respeito.",
    "Comunicar-se de forma assertiva em diferentes contextos.",
    "Planejar o trabalho, cumprir prazos e seguir orientações.",
])

doc.add_heading("Percurso de aprendizagem", level=1)
table = doc.add_table(rows=1, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = "Table Grid"
headers = ["Módulo", "Tema", "Carga horária"]
for i, h in enumerate(headers):
    shade(table.rows[0].cells[i], BLUE)
    p = table.rows[0].cells[i].paragraphs[0]
    font(p.add_run(h), color=WHITE, bold=True, size=10)
set_repeat_table_header(table.rows[0])
for row in [
    ("1", "Leitura e Compreensão Textual", "15h"),
    ("2", "Produção Textual", "12h"),
    ("3", "Gramática Aplicada", "12h"),
    ("4", "Comunicação Oral", "12h"),
    ("5", "Autonomia e Estratégias de Estudo", "12h"),
]:
    cells = table.add_row().cells
    for i, value in enumerate(row):
        font(cells[i].paragraphs[0].add_run(value), size=10)
set_table_widths(table, [0.8, 4.5, 1.2])

# Módulo 1
add_module_intro(doc, 1, "Leitura e Compreensão Textual", "15 horas",
                 "Compreender informações explícitas e implícitas, identificar tema, ideia principal, finalidade, argumentos e confiabilidade das fontes.")
add_lesson(doc, "1.1 Leitura ativa: antes, durante e depois",
           "Ler ativamente significa conversar com o texto. Antes da leitura, observe título, imagens e formato. Durante, destaque palavras-chave e formule perguntas. Depois, resuma o conteúdo e confira se suas previsões estavam corretas.",
           "Ao ver o título 'Celular na escola: ferramenta ou distração?', você pode prever que o texto apresentará vantagens, riscos ou opiniões diferentes.",
           ("Aplicar as três etapas da leitura ativa.", ["Escolha uma notícia curta indicada pelo professor.", "Antes de ler, escreva duas previsões.", "Durante a leitura, destaque cinco palavras-chave.", "Depois, escreva um resumo de três frases."], "Ficha de leitura ativa preenchida.", "40 minutos"))
add_lesson(doc, "1.2 Tema, assunto e ideia principal",
           "O assunto é a área geral abordada. O tema é o recorte específico. A ideia principal é a mensagem mais importante que o autor desenvolve sobre o tema. Para encontrá-la, pergunte: 'o que o texto quer que eu compreenda?'.",
           "Assunto: tecnologia. Tema: uso responsável das redes sociais. Ideia principal: verificar informações antes de compartilhar reduz a circulação de notícias falsas.",
           ("Diferenciar assunto, tema e ideia principal.", ["Leia o texto fornecido pelo professor.", "Escreva o assunto em uma palavra ou expressão.", "Formule o tema em uma frase curta.", "Explique a ideia principal com suas palavras e cite uma evidência do texto."], "Quadro com assunto, tema, ideia principal e evidência.", "45 minutos"))
add_lesson(doc, "1.3 Informações explícitas e inferências",
           "Informações explícitas aparecem diretamente no texto. Inferências são conclusões construídas ao unir pistas do texto com conhecimentos do leitor. Uma boa inferência sempre pode ser justificada por evidências.",
           "Texto: 'Lucas entrou pingando e deixou o guarda-chuva na porta.' Explícito: Lucas estava molhado. Inferência provável: estava chovendo.",
           ("Distinguir fato escrito de conclusão inferida.", ["Leia seis frases preparadas pelo professor.", "Marque E para informação explícita e I para inferência.", "Sublinhe a pista que sustenta cada inferência.", "Compare as respostas com um colega e ajuste o que for necessário."], "Lista classificada e justificada.", "35 minutos"))
add_lesson(doc, "1.4 Intenção comunicativa e gênero textual",
           "Todo texto é produzido com uma finalidade: informar, orientar, convencer, divertir, emocionar ou solicitar algo. O gênero textual organiza essa intenção em uma forma reconhecível, como notícia, receita, tutorial, propaganda, bilhete ou currículo.",
           "Uma campanha de vacinação pode combinar informação ('datas e locais') e convencimento ('proteja você e sua comunidade').",
           ("Reconhecer finalidade e escolhas do autor.", ["Observe quatro textos de gêneros diferentes.", "Identifique o público de cada texto.", "Registre a intenção principal.", "Destaque uma escolha de linguagem que ajuda a cumprir essa intenção."], "Tabela comparativa de gêneros e intenções.", "50 minutos"))
add_lesson(doc, "1.5 Fato, opinião e argumento",
           "Fato é uma informação que pode ser verificada. Opinião apresenta um ponto de vista. Argumento é a razão usada para sustentar uma opinião. Ler criticamente exige separar essas três funções.",
           "Fato: a biblioteca abre às 8h. Opinião: a biblioteca é o melhor espaço da escola. Argumento: ela oferece silêncio, livros e computadores para pesquisa.",
           ("Avaliar argumentos com base em evidências.", ["Leia um pequeno texto de opinião.", "Circule dois fatos e sublinhe duas opiniões.", "Numere os argumentos usados pelo autor.", "Escreva se os argumentos são convincentes e explique por quê."], "Análise crítica de um texto de opinião.", "45 minutos"))
add_lesson(doc, "1.6 Fontes confiáveis e leitura digital",
           "Na internet, a aparência profissional não garante confiança. Verifique autoria, data, instituição responsável, evidências, endereço do site e confirmação em outras fontes. Desconfie de títulos alarmistas e pedidos para compartilhar rapidamente.",
           "Se uma postagem afirma uma descoberta científica, procure a instituição de pesquisa citada e confirme a informação em outra fonte reconhecida.",
           ("Aplicar critérios de confiabilidade.", ["Em grupo, analisem duas páginas indicadas pelo professor.", "Identifiquem autor, data, objetivo e evidências.", "Confirmem uma informação em outra fonte.", "Atribuam uma nota de confiança de 1 a 5 e apresentem a justificativa."], "Checklist de confiabilidade e apresentação breve.", "60 minutos"))
doc.add_heading("Desafio integrador do Módulo 1", level=1)
add_activity(doc, "Dossiê de leitura crítica", "Integrar estratégias de compreensão e análise crítica.",
             ["Escolha um tema atual relacionado à escola ou à juventude.", "Reúna dois textos de gêneros diferentes sobre o tema.", "Identifique tema, ideias principais, fatos, opiniões e intenções.", "Compare as fontes e escreva uma conclusão de um parágrafo."],
             "Dossiê com dois textos analisados e conclusão.", "2 horas")

# Módulo 2
add_module_intro(doc, 2, "Produção Textual", "12 horas",
                 "Planejar, escrever, revisar e reescrever textos claros, coerentes e adequados ao gênero, ao público e à finalidade.")
add_lesson(doc, "2.1 Planejamento: escrever começa antes do texto",
           "Um texto melhora quando o autor define objetivo, leitor, gênero e ideias antes de começar. Mapas mentais, listas e perguntas ajudam a organizar o pensamento.",
           "Para escrever um e-mail solicitando uma visita técnica: objetivo = solicitar; leitor = responsável pela empresa; informações = turma, data, quantidade de alunos e contato.",
           ("Planejar um texto antes da escrita.", ["Escolha uma situação proposta pelo professor.", "Defina quem lerá e qual resultado você espera.", "Liste cinco informações necessárias.", "Organize-as em começo, desenvolvimento e final."], "Plano de texto em tópicos.", "35 minutos"))
add_lesson(doc, "2.2 Parágrafo e organização das ideias",
           "Um parágrafo desenvolve uma ideia central. A frase inicial apresenta o foco; as seguintes explicam, exemplificam ou justificam; a última pode concluir ou fazer ligação com o próximo parágrafo.",
           "Ideia central: manter uma rotina de estudos reduz o acúmulo de tarefas. Desenvolvimento: horário, prioridades e intervalos tornam o trabalho mais organizado.",
           ("Construir parágrafos com unidade temática.", ["Receba quatro frases embaralhadas.", "Organize-as para formar um parágrafo coerente.", "Crie uma frase inicial adequada.", "Escreva um novo parágrafo seguindo a mesma estrutura."], "Dois parágrafos organizados.", "45 minutos"))
add_lesson(doc, "2.3 Coesão: conectando partes do texto",
           "Coesão é a ligação visível entre palavras, frases e parágrafos. Pronomes, sinônimos e conectivos evitam repetições e mostram relações como causa, oposição, sequência e conclusão.",
           "Causa: porque. Consequência: por isso. Oposição: porém. Adição: além disso. Conclusão: portanto.",
           ("Usar conectivos de maneira adequada.", ["Complete dez lacunas com conectivos.", "Explique a relação criada em cada frase.", "Reescreva um parágrafo repetitivo usando pronomes e sinônimos.", "Troque o texto com um colega para revisão."], "Exercícios e parágrafo reescrito.", "45 minutos"))
add_lesson(doc, "2.4 Coerência e progressão temática",
           "Coerência é o sentido global do texto. As informações devem combinar entre si, seguir uma ordem compreensível e permanecer ligadas ao tema. Contradições e mudanças bruscas prejudicam a leitura.",
           "Em um relato, a sequência temporal deve permanecer clara. Se o texto afirma que o evento terminou às 10h, não pode depois dizer que uma atividade começou às 11h e ocorreu antes do encerramento.",
           ("Identificar e corrigir problemas de sentido.", ["Leia um texto com cinco problemas de coerência.", "Marque contradições e informações fora do tema.", "Reordene os trechos quando necessário.", "Produza uma versão corrigida."], "Texto revisado com justificativas.", "50 minutos"))
add_lesson(doc, "2.5 Gêneros do cotidiano e do mundo do trabalho",
           "Mensagens, e-mails, avisos, relatos e currículos exigem níveis diferentes de formalidade. A linguagem deve respeitar o leitor e apresentar as informações essenciais sem excesso nem abreviações inadequadas.",
           "Em vez de 'Oi, manda aí o horário', escreva: 'Olá, professora. Poderia informar o horário da atividade de amanhã? Obrigado.'",
           ("Adaptar linguagem e formato ao contexto.", ["Transforme uma mensagem informal em e-mail profissional.", "Inclua assunto, saudação, pedido claro e despedida.", "Revise ortografia e pontuação.", "Leia o texto como se fosse o destinatário e faça o último ajuste."], "E-mail profissional completo.", "45 minutos"))
add_lesson(doc, "2.6 Revisão e reescrita",
           "Revisar não é apenas procurar erros ortográficos. É verificar se o texto cumpre o objetivo, se as ideias estão organizadas, se há ligação entre as partes e se a linguagem está adequada. Reescrever é melhorar com base nessa análise.",
           "Checklist: o título combina com o texto? Cada parágrafo tem foco? Há informações repetidas? Os conectivos estão adequados? A pontuação ajuda a leitura?",
           ("Utilizar uma rotina de revisão em camadas.", ["Escolha um texto produzido anteriormente.", "Primeira leitura: confira sentido e objetivo.", "Segunda leitura: observe organização e conectivos.", "Terceira leitura: corrija pontuação, concordância e ortografia."], "Primeira versão, checklist e versão final.", "60 minutos"))
doc.add_heading("Projeto de escrita do Módulo 2", level=1)
add_activity(doc, "Artigo de opinião: uma melhoria para a escola", "Defender um ponto de vista com clareza e argumentos.",
             ["Escolha uma melhoria possível para a escola ou comunidade.", "Formule sua opinião e dois argumentos.", "Escreva introdução, desenvolvimento e conclusão.", "Revise com o checklist e produza a versão final."],
             "Artigo de opinião de 20 a 30 linhas.", "2 horas")

# Módulo 3
add_module_intro(doc, 3, "Gramática Aplicada", "12 horas",
                 "Usar pontuação, concordância e recursos linguísticos para tornar textos mais claros, precisos e adequados às situações reais de comunicação.")
add_lesson(doc, "3.1 Pontuação e construção de sentido",
           "A pontuação organiza a leitura e pode alterar o sentido. O ponto final encerra uma ideia; a vírgula separa termos e orações; dois-pontos anunciam explicação ou lista; ponto de interrogação marca pergunta; exclamação expressa intensidade.",
           "Compare: 'Não espere.' e 'Não, espere.' Uma vírgula muda completamente a orientação.",
           ("Perceber o efeito da pontuação.", ["Pontue um diálogo curto que está sem sinais.", "Leia as versões em voz alta.", "Explique duas mudanças de sentido.", "Crie um exemplo em que a vírgula altere a mensagem."], "Diálogo pontuado e exemplos comentados.", "45 minutos"))
add_lesson(doc, "3.2 Usos essenciais da vírgula",
           "A vírgula pode separar itens de uma enumeração, marcar vocativo, isolar explicações e indicar deslocamento de expressões. Em geral, não se separa sujeito de verbo nem verbo de complemento.",
           "Correto: 'Pedro, entregue o relatório.' Incorreto: 'Os estudantes da turma, terminaram a atividade.'",
           ("Aplicar regras frequentes de vírgula.", ["Corrija dez frases.", "Ao lado de cada vírgula, indique a função.", "Identifique duas vírgulas indevidas.", "Escreva uma frase com enumeração e outra com vocativo."], "Lista corrigida e classificada.", "45 minutos"))
add_lesson(doc, "3.3 Concordância verbal",
           "Na concordância verbal, o verbo combina com o núcleo do sujeito em número e pessoa. Para acertar, identifique quem pratica ou recebe a ação e localize a palavra principal do sujeito.",
           "'Os resultados da pesquisa mostram mudanças.' O núcleo do sujeito é 'resultados', por isso o verbo fica no plural.",
           ("Relacionar sujeito e verbo.", ["Sublinhe o sujeito de oito frases.", "Circule o núcleo do sujeito.", "Escolha a forma verbal correta.", "Reescreva três frases mudando o sujeito do singular para o plural."], "Exercícios de concordância verbal.", "50 minutos"))
add_lesson(doc, "3.4 Concordância nominal",
           "Artigos, adjetivos, numerais e pronomes concordam com o substantivo ao qual se referem. Observar gênero e número ajuda a evitar construções que prejudicam a clareza.",
           "'As duas novas propostas foram aprovadas.' As palavras 'as', 'duas' e 'novas' acompanham 'propostas'.",
           ("Aplicar concordância dentro do grupo nominal.", ["Complete frases com as formas adequadas.", "Ligue cada adjetivo ao substantivo correspondente.", "Corrija um anúncio com erros de concordância.", "Crie duas frases corretas no plural."], "Anúncio revisado e frases autorais.", "45 minutos"))
add_lesson(doc, "3.5 Classes de palavras em uso",
           "Substantivos nomeiam; adjetivos caracterizam; verbos indicam ações, estados ou fenômenos; advérbios modificam sentidos; pronomes retomam termos; conectivos relacionam ideias. O importante é perceber a função no texto.",
           "Em 'A equipe respondeu rapidamente porque estava preparada', 'rapidamente' modifica 'respondeu' e 'porque' apresenta a causa.",
           ("Analisar funções das palavras no contexto.", ["Leia um parágrafo curto.", "Marque substantivos, verbos, adjetivos e conectivos com cores diferentes.", "Explique o efeito de dois adjetivos.", "Troque um advérbio e descreva a mudança de sentido."], "Parágrafo analisado e comentado.", "45 minutos"))
add_lesson(doc, "3.6 Clareza, ambiguidade e escolha vocabular",
           "Ambiguidade ocorre quando uma frase permite mais de uma interpretação. Pronomes sem referência clara, ordem confusa e palavras vagas podem causar o problema. A solução é nomear os elementos e reorganizar a frase.",
           "Ambígua: 'Ana falou com Júlia sobre seu projeto.' Clara: 'Ana falou com Júlia sobre o projeto de Júlia.'",
           ("Reescrever frases para eliminar ambiguidades.", ["Identifique duas interpretações possíveis em cinco frases.", "Sublinhe a causa da ambiguidade.", "Reescreva cada frase com sentido claro.", "Troque expressões vagas por palavras mais precisas."], "Banco de frases revisadas.", "50 minutos"))
doc.add_heading("Oficina do Módulo 3", level=1)
add_activity(doc, "Clínica de textos", "Aplicar gramática como ferramenta de comunicação.",
             ["Em dupla, recebam um texto com problemas de pontuação, concordância e clareza.", "Façam uma revisão por categorias.", "Registrem cada correção e sua justificativa.", "Apresentem duas mudanças que mais melhoraram o texto."],
             "Texto corrigido e relatório breve de revisão.", "2 horas")

# Módulo 4
add_module_intro(doc, 4, "Comunicação Oral", "12 horas",
                 "Planejar falas, escutar ativamente, argumentar com respeito e apresentar ideias de modo claro, assertivo e adequado ao público.")
add_lesson(doc, "4.1 Comunicação verbal, não verbal e paraverbal",
           "A comunicação oral envolve palavras, postura, gestos, olhar, volume, ritmo e entonação. Esses elementos devem reforçar a mensagem. Uma fala correta pode parecer insegura quando é muito baixa ou rápida.",
           "Dizer 'estou preparado' olhando para o chão e quase sem voz transmite uma mensagem diferente das palavras.",
           ("Observar como corpo e voz influenciam a mensagem.", ["Em trio, escolha uma frase neutra.", "Diga a frase com três intenções diferentes.", "Os colegas identificam a intenção e apontam pistas.", "Registre quais elementos ajudaram mais."], "Ficha de observação da comunicação.", "40 minutos"))
add_lesson(doc, "4.2 Escuta ativa e empatia",
           "Escutar ativamente é dar atenção, não interromper, fazer perguntas para compreender e confirmar o que foi entendido. Empatia é considerar a perspectiva do outro sem precisar concordar com tudo.",
           "Uma confirmação útil é: 'Entendi que o prazo preocupa você porque ainda faltam informações. É isso?'.",
           ("Praticar escuta e confirmação de entendimento.", ["Formem duplas.", "Uma pessoa fala por dois minutos sobre um desafio de estudo.", "A outra escuta sem interromper e depois resume.", "Troquem os papéis e avaliem a qualidade da escuta."], "Autoavaliação de escuta ativa.", "45 minutos"))
add_lesson(doc, "4.3 Assertividade e respeito",
           "Ser assertivo é comunicar necessidades, limites e opiniões com clareza e respeito. A fala passiva esconde o que a pessoa precisa; a agressiva desrespeita; a assertiva busca solução.",
           "Estrutura: fato + efeito + necessidade + pedido. 'Quando há conversa durante a apresentação, perco a concentração. Preciso de silêncio. Podemos combinar isso?'.",
           ("Transformar falas passivas ou agressivas em assertivas.", ["Leia seis situações de conflito cotidiano.", "Identifique o estilo de comunicação.", "Reescreva as falas usando fato, efeito, necessidade e pedido.", "Encene uma situação com um colega."], "Roteiro e encenação assertiva.", "50 minutos"))
add_lesson(doc, "4.4 Argumentação em debates",
           "Um argumento oral precisa de uma afirmação clara, uma razão e uma evidência ou exemplo. Em debates, responda à ideia e não ataque a pessoa. Reconhecer um ponto válido do outro fortalece o diálogo.",
           "Afirmação: a escola deve ampliar a coleta seletiva. Razão: muitos resíduos recicláveis ainda vão para o lixo comum. Evidência: observação realizada durante uma semana.",
           ("Construir e apresentar argumentos respeitosos.", ["Em grupo, escolham um tema proposto.", "Preparem uma afirmação, duas razões e exemplos.", "Antecipem um contra-argumento.", "Realizem um debate curto com regras de tempo e escuta."], "Mapa de argumentos e participação no debate.", "60 minutos"))
add_lesson(doc, "4.5 Planejamento de apresentações",
           "Uma apresentação eficiente tem abertura que situa o tema, desenvolvimento organizado e encerramento que retoma a mensagem principal. Slides apoiam a fala; não devem substituir o apresentador.",
           "Roteiro de três minutos: 20 segundos para abertura, 2 minutos para dois pontos principais e 40 segundos para síntese e conclusão.",
           ("Organizar uma apresentação curta.", ["Escolha um tema que domina.", "Defina a mensagem principal em uma frase.", "Crie roteiro com abertura, dois pontos e conclusão.", "Prepare no máximo três apoios visuais simples."], "Roteiro e apoios visuais.", "60 minutos"))
add_lesson(doc, "4.6 Voz, postura e controle do nervosismo",
           "O nervosismo é uma reação comum. Preparação, respiração, ensaio e foco na mensagem reduzem a ansiedade. Falar um pouco mais devagar e fazer pausas ajuda o público e o apresentador.",
           "Antes de começar: apoie os pés no chão, inspire lentamente, olhe para três pontos da sala e diga a primeira frase já ensaiada.",
           ("Aplicar técnicas de preparação e autoavaliação.", ["Faça dois ciclos de respiração lenta.", "Apresente seu roteiro por dois minutos.", "Um colega observa clareza, volume, ritmo e postura.", "Repita a fala aplicando uma melhoria."], "Ficha de feedback e segunda apresentação.", "50 minutos"))
doc.add_heading("Apresentação final do Módulo 4", level=1)
add_activity(doc, "Ideia que transforma", "Apresentar uma proposta de melhoria com argumentos e comunicação assertiva.",
             ["Escolha um problema próximo da realidade da turma.", "Apresente a proposta, os benefícios e uma ação inicial.", "Use fala de três a cinco minutos e apoio visual opcional.", "Responda a uma pergunta do público e registre o feedback."],
             "Apresentação oral e ficha de autoavaliação.", "2 horas")

# Módulo 5
add_module_intro(doc, 5, "Autonomia e Estratégias de Estudo", "12 horas",
                 "Planejar estudos, administrar tempo, tomar notas, pesquisar com responsabilidade, acompanhar o próprio progresso e colaborar em grupo.")
add_lesson(doc, "5.1 Metas e planejamento semanal",
           "Metas eficazes são específicas, possíveis e ligadas a um prazo. Um planejamento realista considera compromissos, energia disponível, prioridades e intervalos.",
           "Em vez de 'estudar português', use: 'na terça, das 18h às 18h40, revisar vírgula e resolver cinco exercícios'.",
           ("Transformar intenções em ações planejadas.", ["Liste três responsabilidades da semana.", "Defina uma meta específica de aprendizagem.", "Divida a meta em tarefas de até 40 minutos.", "Reserve horários e um momento de revisão."], "Plano semanal de estudos.", "45 minutos"))
add_lesson(doc, "5.2 Gestão do tempo e foco",
           "Foco depende do ambiente e da maneira como a tarefa é dividida. Retire distrações, escolha uma pequena meta, trabalhe por um período definido e faça uma pausa curta antes do próximo ciclo.",
           "Ciclo possível: 25 minutos de concentração, 5 minutos de pausa e mais 25 minutos para concluir ou revisar.",
           ("Testar um ciclo de estudo focado.", ["Escolha uma tarefa curta.", "Prepare o ambiente e silencie notificações.", "Trabalhe por 25 minutos.", "Registre o que concluiu e qual distração apareceu."], "Registro de foco e plano de melhoria.", "40 minutos"))
add_lesson(doc, "5.3 Anotações, resumos e mapas mentais",
           "Boas anotações selecionam ideias, não copiam tudo. Um resumo preserva as informações essenciais com palavras do estudante. O mapa mental organiza conceitos por relações e palavras-chave.",
           "Após uma explicação sobre coesão, o centro do mapa pode ser 'coesão' e os ramos: pronomes, sinônimos, conectivos e relações de sentido.",
           ("Comparar técnicas de registro.", ["Leia um texto de uma página.", "Faça anotações em tópicos.", "Produza um resumo de cinco linhas.", "Crie um mapa mental e escolha o formato que mais ajudou."], "Três registros e justificativa da escolha.", "55 minutos"))
add_lesson(doc, "5.4 Pesquisa e uso ético das informações",
           "Pesquisar exige formular uma pergunta, escolher palavras-chave, avaliar fontes e registrar de onde vieram as informações. Copiar sem indicar a fonte é plágio; aprender implica compreender e reescrever com autoria.",
           "Anote: autor ou instituição, título, endereço, data de publicação e data de acesso. Use aspas apenas para uma citação curta e necessária.",
           ("Realizar pesquisa breve com responsabilidade.", ["Formule uma pergunta sobre um tema do curso.", "Pesquise em duas fontes confiáveis.", "Registre os dados das fontes.", "Escreva um parágrafo com suas palavras e indique as fontes."], "Parágrafo informativo e referências.", "60 minutos"))
add_lesson(doc, "5.5 Aprender com erros e feedback",
           "Feedback mostra o que funcionou e o que pode melhorar. Para usá-lo, separe a tarefa da sua identidade: um texto precisa de ajustes, mas isso não define sua capacidade. Escolha uma mudança concreta e tente novamente.",
           "Feedback: 'a conclusão apresenta uma ideia nova'. Ação: mover a ideia para o desenvolvimento e concluir retomando a tese.",
           ("Transformar feedback em plano de ação.", ["Leia um feedback recebido em atividade anterior.", "Identifique um ponto forte e uma melhoria.", "Escreva uma ação concreta.", "Refaça o trecho e compare as versões."], "Plano de melhoria e trecho reescrito.", "45 minutos"))
add_lesson(doc, "5.6 Colaboração, responsabilidades e prazos",
           "Trabalho em grupo exige objetivo comum, divisão justa, comunicação e acompanhamento. Cada integrante precisa saber o que fará, até quando e como a entrega será integrada.",
           "Quadro do grupo: tarefa, responsável, prazo, status e ajuda necessária. Reuniões curtas evitam surpresas no final.",
           ("Planejar uma entrega colaborativa.", ["Formem grupos e escolham uma produção final.", "Dividam o trabalho em tarefas pequenas.", "Definam responsáveis e prazos.", "Façam uma reunião de cinco minutos para verificar o andamento."], "Quadro de responsabilidades do grupo.", "50 minutos"))
doc.add_heading("Projeto integrador final", level=1)
add_activity(doc, "Guia de comunicação para estudantes", "Integrar leitura, escrita, gramática, oralidade e autonomia.",
             ["Em grupo, escolham um desafio de comunicação vivido por estudantes.", "Pesquisem informações confiáveis e organizem as ideias.", "Produzam um guia de duas páginas com orientações claras.", "Revisem o texto e apresentem o guia à turma em até cinco minutos."],
             "Guia escrito, referências, apresentação e autoavaliação.", "3 horas")

# Instrumentos finais
doc.add_page_break()
doc.add_heading("Checklists de apoio", level=1)
doc.add_heading("Checklist de leitura", level=2)
add_bullets(doc, ["Observei título, fonte, imagens e gênero.", "Identifiquei tema e ideia principal.", "Diferenciei fatos, opiniões e inferências.", "Localizei evidências para minhas conclusões.", "Consegui resumir o texto com minhas palavras."])
doc.add_heading("Checklist de escrita", level=2)
add_bullets(doc, ["Defini objetivo, público e gênero.", "Planejei começo, desenvolvimento e final.", "Cada parágrafo desenvolve uma ideia central.", "Usei conectivos e evitei repetições desnecessárias.", "Revisei coerência, pontuação, concordância e ortografia."])
doc.add_heading("Checklist de apresentação oral", level=2)
add_bullets(doc, ["Minha mensagem principal está clara.", "Organizei abertura, desenvolvimento e conclusão.", "Ensaiar ajudou a controlar o tempo.", "Usei volume, ritmo, postura e olhar adequados.", "Escutei perguntas e respondi com respeito."])

doc.add_heading("Autoavaliação final", level=1)
table = doc.add_table(rows=1, cols=4)
table.style = "Table Grid"
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(["Competência", "Ainda preciso de ajuda", "Estou avançando", "Faço com autonomia"]):
    shade(table.rows[0].cells[i], BLUE)
    font(table.rows[0].cells[i].paragraphs[0].add_run(h), color=WHITE, bold=True, size=9)
set_repeat_table_header(table.rows[0])
for item in ["Compreender textos", "Identificar tema e ideia principal", "Escrever parágrafos coerentes", "Revisar meus textos", "Usar pontuação e concordância", "Apresentar ideias oralmente", "Planejar meus estudos", "Colaborar e cumprir prazos"]:
    cells = table.add_row().cells
    font(cells[0].paragraphs[0].add_run(item), size=9.5)
    for c in cells[1:]:
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        font(c.paragraphs[0].add_run("☐"), size=12, color=BLUE)
set_table_widths(table, [2.9, 1.2, 1.2, 1.2])

doc.add_heading("Gabarito orientativo", level=1)
add_para(doc, "Muitas atividades possuem respostas pessoais. O gabarito abaixo apresenta critérios para correção e exemplos possíveis, sem substituir a análise do contexto.", italic=True, color=GRAY)
for title, answer in [
    ("Leitura ativa", "As previsões devem estar relacionadas ao título; as palavras-chave precisam representar as ideias centrais; o resumo deve ser fiel e não incluir opinião pessoal."),
    ("Tema e ideia principal", "O assunto é amplo, o tema é o recorte e a ideia principal expressa a mensagem central sustentada pelo texto."),
    ("Inferência", "Toda inferência precisa ser ligada a uma pista do texto. Respostas possíveis variam conforme os conhecimentos prévios, mas não podem contradizer as evidências."),
    ("Fato, opinião e argumento", "Fatos são verificáveis; opiniões expressam julgamentos; argumentos justificam um ponto de vista com razões ou evidências."),
    ("Coesão", "Conectivos devem respeitar a relação: porque = causa; porém = oposição; além disso = adição; portanto = conclusão."),
    ("Vírgula", "É adequada em enumeração, vocativo, explicação intercalada e expressão deslocada. Não deve separar sujeito e verbo."),
    ("Concordância", "O verbo concorda com o núcleo do sujeito; determinantes e adjetivos concordam com o substantivo a que se referem."),
    ("Assertividade", "Uma resposta adequada descreve o fato sem acusação, indica o efeito ou necessidade e faz um pedido possível e respeitoso."),
    ("Apresentação", "Avaliar clareza da mensagem, organização, evidências, adequação ao tempo, voz, postura, contato com o público e resposta às perguntas."),
    ("Projeto integrador", "O guia deve apresentar problema definido, informações confiáveis, orientações práticas, texto revisado, fontes e participação equilibrada do grupo."),
]:
    doc.add_heading(title, level=2)
    add_para(doc, answer)

doc.add_heading("Glossário", level=1)
for term, definition in [
    ("Argumento", "Razão ou evidência usada para sustentar uma opinião."),
    ("Coerência", "Relação lógica e sentido global de um texto."),
    ("Coesão", "Recursos que conectam palavras, frases e parágrafos."),
    ("Explícito", "Aquilo que está declarado diretamente."),
    ("Gênero textual", "Forma social de texto ligada a uma finalidade, como notícia, e-mail ou tutorial."),
    ("Inferência", "Conclusão construída a partir de pistas e conhecimentos prévios."),
    ("Intenção comunicativa", "Objetivo que orienta a produção de uma mensagem."),
    ("Parágrafo", "Unidade textual que desenvolve uma ideia central."),
    ("Reescrita", "Produção de nova versão para melhorar conteúdo e forma."),
    ("Tese", "Ponto de vista central defendido em um texto argumentativo."),
]:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    font(p.add_run(f"{term}: "), bold=True, color=BLUE)
    font(p.add_run(definition))

doc.add_heading("Referências pedagógicas sugeridas", level=1)
add_bullets(doc, [
    "BRASIL. Base Nacional Comum Curricular. Brasília: MEC, 2018.",
    "Dicionários e gramáticas escolares atualizados da língua portuguesa.",
    "Textos jornalísticos, campanhas, tutoriais e gêneros digitais selecionados pelo docente conforme a turma.",
])

doc.core_properties.title = "Apostila de Reforço de Linguagens"
doc.core_properties.subject = "Leitura, escrita, gramática, comunicação oral e estratégias de estudo"
doc.core_properties.author = "SENAI | Rio do Sul Mais Tech"
doc.core_properties.keywords = "reforço, linguagens, leitura, escrita, oralidade, ensino fundamental"
doc.core_properties.comments = "Material baseado na Ementa de Reforço de Linguagens, carga horária de 63 horas."

OUTPUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUTPUT)
print(OUTPUT.resolve())
