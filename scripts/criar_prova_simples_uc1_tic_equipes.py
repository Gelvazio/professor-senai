from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"C:\fontes\professor-senai")
OUTPUT = ROOT / "sistema" / "INTRODUCAO_A_TECNOLOGIA_DA_INFORMACAO_E_COMUNICACAO" / "AVALIACOES_CRIADAS" / "PROVA_PRATICA" / "PROVA_SIMPLES_UC1_TIC_EQUIPES_2H.docx"

BLUE = "1F4E78"
LIGHT_BLUE = "D9EAF7"
LIGHT_GRAY = "F2F4F7"
WHITE = "FFFFFF"
INK = RGBColor(31, 41, 55)


def set_cell_fill(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), color)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_widths(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths_dxa = [round(width * 1440) for width in widths]
    total_dxa = sum(widths_dxa)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width_dxa in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width_dxa))
        grid.append(grid_col)
    for row in table.rows:
        for idx, width_dxa in enumerate(widths_dxa):
            cell = row.cells[idx]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width_dxa))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def style_run(run, size=10.5, bold=False, color=INK, font="Calibri"):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = color


def add_paragraph(doc, text="", bold_prefix=None, after=5, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.08
    if align is not None:
        p.alignment = align
    if bold_prefix and text.startswith(bold_prefix):
        style_run(p.add_run(bold_prefix), bold=True)
        style_run(p.add_run(text[len(bold_prefix):]))
    else:
        style_run(p.add_run(text))
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.38)
    p.paragraph_format.first_line_indent = Inches(-0.19)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.08
    style_run(p.add_run(text), size=10.5)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.space_before = Pt(10 if level == 1 else 7)
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run(text)
    style_run(run, size=15 if level == 1 else 12, bold=True, color=RGBColor(31, 78, 120))
    return p


def add_info_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    set_table_widths(table, [1.55, 4.95])
    for idx, (label, value) in enumerate(rows):
        set_cell_fill(table.cell(idx, 0), LIGHT_BLUE)
        p1 = table.cell(idx, 0).paragraphs[0]
        p2 = table.cell(idx, 1).paragraphs[0]
        style_run(p1.add_run(label), bold=True, color=RGBColor(31, 78, 120))
        style_run(p2.add_run(value))
    return table


def add_rubric(doc):
    rows = [
        ("Organização e entrega — automática", "Links acessíveis, arquivos Google nativos, nomes padronizados e identificação da equipe.", "1,0"),
        ("Google Docs — automática", "Estrutura, conteúdos obrigatórios, termos técnicos, práticas de segurança e fonte.", "2,0"),
        ("Google Sheets — automática", "Cabeçalhos, seis registros, fórmula, filtro e formatação básica.", "2,0"),
        ("Google Slides — automática", "Quatro slides, identificação, conteúdos, imagem e quantidade adequada de texto.", "1,5"),
        ("Segurança e interpretação — automática", "Presença dos conceitos mínimos definidos no Anexo I.", "0,5"),
        ("Clareza e linguagem — docente", "Texto claro, formal, organizado e adequado ao ambiente profissional.", "1,0"),
        ("Correção técnica — docente", "Explicações corretas, coerentes e adaptadas ao público.", "1,0"),
        ("Apresentação e cooperação — docente", "Participação dos dois integrantes, domínio e colaboração.", "1,0"),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    set_table_widths(table, [1.65, 4.25, 0.60])
    for idx, text in enumerate(("Critério", "Evidência esperada", "Valor")):
        set_cell_fill(table.cell(0, idx), BLUE)
        p = table.cell(0, idx).paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style_run(p.add_run(text), bold=True, color=RGBColor(255, 255, 255))
    for criterion, evidence, score in rows:
        cells = table.add_row().cells
        for idx, text in enumerate((criterion, evidence, score)):
            if len(table.rows) % 2 == 0:
                set_cell_fill(cells[idx], LIGHT_GRAY)
            p = cells[idx].paragraphs[0]
            if idx == 2:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            style_run(p.add_run(text), size=9.5, bold=(idx == 0))
            set_cell_margins(cells[idx])
    return table


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.65)
section.bottom_margin = Inches(0.65)
section.left_margin = Inches(0.75)
section.right_margin = Inches(0.75)
section.header_distance = Inches(0.30)
section.footer_distance = Inches(0.30)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(10.5)
normal.paragraph_format.space_after = Pt(5)
normal.paragraph_format.line_spacing = 1.08

header = section.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.CENTER
style_run(header.add_run("SENAI — Introdução à Tecnologia da Informação e Comunicação"), size=8.5, bold=True, color=RGBColor(90, 100, 115))

footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
style_run(footer.add_run("Avaliação prática em equipe • duração máxima: 2 horas"), size=8, color=RGBColor(90, 100, 115))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(2)
style_run(p.add_run("AVALIAÇÃO PRÁTICA SIMPLES"), size=20, bold=True, color=RGBColor(31, 78, 120))
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(10)
style_run(p.add_run("Missão: organizar e proteger o computador do setor"), size=12.5, bold=True, color=RGBColor(65, 75, 90))

add_info_table(doc, [
    ("Unidade curricular", "Introdução à Tecnologia da Informação e Comunicação"),
    ("Modalidade", "Trabalho em equipe, realizado em sala de aula"),
    ("Duração", "2 horas"),
    ("Estudante 1", "________________________________________________________________"),
    ("Estudante 2", "________________________________________________________________"),
    ("Data / Turma", "Data: ____/____/________    Turma: ______________________________"),
])

add_heading(doc, "Divisão de responsabilidades", 1)
add_paragraph(doc, "Cada integrante terá um foco principal, mas ambos devem revisar, conhecer e saber explicar todas as entregas.")
responsibilities = doc.add_table(rows=1, cols=4)
responsibilities.style = "Table Grid"
set_table_widths(responsibilities, [1.00, 1.60, 1.70, 2.20])
for idx, value in enumerate(("Integrante", "Nome", "Foco principal", "Responsabilidade")):
    set_cell_fill(responsibilities.cell(0, idx), BLUE)
    p = responsibilities.cell(0, idx).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_run(p.add_run(value), size=9, bold=True, color=RGBColor(255, 255, 255))
for row in [
    ("Estudante 1", "________________", "Desafio 1 — Guia", "Liderar o Google Docs e a interpretação do Anexo I."),
    ("Estudante 2", "________________", "Desafio 2 — Inventário", "Liderar o Google Sheets, os registros e a fórmula."),
    ("Toda a equipe", "Ambos", "Desafio 3 — Apresentação", "Produzir os slides, revisar tudo, apresentar e entregar."),
]:
    cells = responsibilities.add_row().cells
    for idx, value in enumerate(row):
        p = cells[idx].paragraphs[0]
        if idx < 3:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style_run(p.add_run(value), size=8.8, bold=(idx == 0))
        set_cell_margins(cells[idx])

add_heading(doc, "Situação-problema", 1)
add_paragraph(doc, "O computador de um pequeno setor da empresa está desorganizado. Os arquivos estão espalhados, não existe rotina de backup, alguns funcionários recebem mensagens suspeitas e ninguém sabe explicar com clareza quais equipamentos e programas são utilizados. A equipe foi escolhida para preparar uma solução simples e orientar o setor.")

add_heading(doc, "Organização do tempo", 1)
time_table = doc.add_table(rows=1, cols=3)
time_table.style = "Table Grid"
set_table_widths(time_table, [1.15, 4.55, 0.80])
for idx, value in enumerate(("Etapa", "Atividade", "Tempo")):
    set_cell_fill(time_table.cell(0, idx), BLUE)
    p = time_table.cell(0, idx).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_run(p.add_run(value), bold=True, color=RGBColor(255, 255, 255))
for row in [
    ("1", "Ler, dividir funções e criar a pasta da equipe", "10 min"),
    ("2", "Produzir as três entregas", "85 min"),
    ("3", "Apresentar e entregar os arquivos", "15 min"),
    ("4", "Margem para revisão e imprevistos", "10 min"),
]:
    cells = time_table.add_row().cells
    for idx, value in enumerate(row):
        p = cells[idx].paragraphs[0]
        if idx != 1:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style_run(p.add_run(value), size=9.5)
        set_cell_margins(cells[idx])

add_heading(doc, "Regras simples", 1)
for item in [
    "O Estudante 1 lidera o Desafio 1; o Estudante 2 lidera o Desafio 2; os dois realizam o Desafio 3.",
    "O foco principal organiza o trabalho, mas os dois integrantes precisam conhecer e revisar todas as entregas.",
    "Produzam os três arquivos nos formatos nativos Google Docs, Google Sheets e Google Slides.",
    "Criem uma pasta com o nome EQUIPE_NOME1_NOME2 e salvem nela os três arquivos.",
    "É permitido consultar a ementa, as anotações de aula e fontes confiáveis na Web.",
    "Toda informação pesquisada deve ter o nome do site ou o link registrado no documento.",
    "Compartilhem os três arquivos com o docente; sem acesso, o corretor não conseguirá avaliá-los.",
    "Ao final, enviem os três links no Formulário de Entrega fornecido pelo docente.",
    "A apresentação final deve durar no máximo 3 minutos.",
]:
    add_bullet(doc, item)

doc.add_page_break()
add_heading(doc, "Entrega 1 — Guia rápido para a equipe", 1)
add_paragraph(doc, "Criem um documento de 1 a 2 páginas no editor de textos com o título “Como organizar e proteger o computador do setor”. O guia deve conter:")
for item in [
    "um parágrafo curto explicando o problema e a solução proposta;",
    "uma explicação simples da diferença entre hardware e software, com 2 exemplos de cada;",
    "a função de CPU, memória RAM, HD ou SSD e um periférico;",
    "4 boas práticas de segurança, incluindo senha forte, backup, navegação segura e cuidado com golpes;",
    "a interpretação do Anexo I, explicando em linguagem acessível os termos backup, malware e autenticação;",
    "uma fonte consultada, identificada pelo nome do site ou link;",
    "formatação organizada: título, subtítulos, parágrafos, marcadores e correção ortográfica.",
]:
    add_bullet(doc, item)
add_paragraph(doc, "Nome obrigatório do arquivo: GUIA_DA_EQUIPE (Google Docs)", bold_prefix="Nome obrigatório do arquivo:")

add_heading(doc, "Entrega 2 — Inventário básico", 1)
add_paragraph(doc, "Criem uma planilha com pelo menos 6 itens do setor. Usem as colunas abaixo:")
inventory = doc.add_table(rows=2, cols=6)
inventory.style = "Table Grid"
set_table_widths(inventory, [0.65, 1.10, 1.55, 0.85, 1.10, 1.25])
headers = ("Código", "Item", "Tipo", "Qtd.", "Estado", "Ação necessária")
for idx, value in enumerate(headers):
    set_cell_fill(inventory.cell(0, idx), BLUE)
    p = inventory.cell(0, idx).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_run(p.add_run(value), size=9, bold=True, color=RGBColor(255, 255, 255))
examples = ("EQ-01", "Monitor", "Hardware", "2", "Bom", "Limpar")
for idx, value in enumerate(examples):
    p = inventory.cell(1, idx).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_run(p.add_run(value), size=9)
add_paragraph(doc, "Use a linha 1 para os cabeçalhos e as linhas 2 a 7 para os seis primeiros registros. Em G2, escreva TOTAL DE ITENS. Em H2, insira uma fórmula que some as quantidades da coluna D. Aplique filtro na linha 1 e destaque os cabeçalhos com uma cor de fundo.")
add_paragraph(doc, "Nome obrigatório do arquivo: INVENTARIO_DA_EQUIPE (Google Sheets)", bold_prefix="Nome obrigatório do arquivo:")

add_heading(doc, "Entrega 3 — Apresentação-relâmpago", 1)
add_paragraph(doc, "Criem uma apresentação com exatamente 4 slides:")
for item in [
    "Slide 1 — título da missão e nomes da equipe;",
    "Slide 2 — hardware e software encontrados no setor;",
    "Slide 3 — três riscos digitais e como evitá-los;",
    "Slide 4 — organização de arquivos, backup e conclusão da equipe.",
]:
    add_bullet(doc, item)
add_paragraph(doc, "Usem pouco texto, letras legíveis e pelo menos uma imagem ou ícone adequado. Os dois integrantes devem participar da apresentação.")
add_paragraph(doc, "Nome obrigatório do arquivo: APRESENTACAO_DA_EQUIPE (Google Slides)", bold_prefix="Nome obrigatório do arquivo:")

doc.add_page_break()
add_heading(doc, "Anexo I — Mensagem técnica", 1)
anexo = doc.add_table(rows=1, cols=1)
anexo.style = "Table Grid"
set_table_widths(anexo, [6.5])
set_cell_fill(anexo.cell(0, 0), LIGHT_BLUE)
p = anexo.cell(0, 0).paragraphs[0]
p.paragraph_format.space_after = Pt(0)
style_run(p.add_run("“Para reduzir o risco de perda de dados, o setor deve realizar backup semanal dos arquivos em local seguro. Os computadores precisam manter o antimalware atualizado e as contas devem utilizar autenticação com senha forte. Mensagens com links desconhecidos não devem ser abertas sem verificação.”"), size=11)

add_heading(doc, "O que explicar no guia", 2)
for item in [
    "Qual é a orientação principal do texto?",
    "O que significam backup, malware e autenticação?",
    "Como você explicaria toda a mensagem a um colega sem conhecimento técnico?",
]:
    add_bullet(doc, item)

add_heading(doc, "Checklist antes de entregar", 1)
for item in [
    "A pasta está com o nome correto e contém os três arquivos.",
    "Os arquivos estão nos formatos nativos Google Docs, Google Sheets e Google Slides.",
    "O guia tem os itens pedidos e registra uma fonte consultada.",
    "A planilha tem 6 ou mais registros, filtro, cabeçalhos na linha 1 e fórmula em H2.",
    "A apresentação tem exatamente 4 slides e está legível.",
    "Os três arquivos foram compartilhados com o docente e seus links foram enviados no formulário.",
    "Os dois integrantes conseguem explicar o trabalho realizado.",
]:
    add_bullet(doc, "☐ " + item)

add_heading(doc, "Critérios de avaliação — 10 pontos", 1)
add_paragraph(doc, "A nota é composta por 7,0 pontos de verificação automática e 3,0 pontos de avaliação do docente. A correção automática confirma requisitos objetivos; a qualidade e a correção das explicações continuam sob responsabilidade do professor.")
add_rubric(doc)

add_paragraph(doc, "Observação do docente: __________________________________________________________________________________", after=8)
add_paragraph(doc, "Nota: ______ / 10,0     Assinatura da equipe: _____________________________________________", after=0)

OUTPUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUTPUT)
print(OUTPUT)
