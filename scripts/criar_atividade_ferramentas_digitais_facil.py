from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from criar_apostila_ferramentas_digitais import (
    BLUE,
    GRAY,
    LIGHT_BLUE,
    ORANGE,
    add_body,
    add_bullets,
    add_callout,
    add_two_column_table,
    configure_document,
    set_cell_margins,
    set_cell_shading,
    style_run,
)


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "sistema" / "FICHA-PRODUTO-MAIS-TECH" / "INTRODUCAO_COMUNICACAO_ORAL_ESCRITA" / "AtividadesFerramentasDigitaisFacil.docx"


def line(doc, label=""):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(7)
    if label:
        style_run(p.add_run(label + " "), size=10.5, color=GRAY, bold=True)
    style_run(p.add_run("________________________________________________________________"), size=10, color="A0A0A0")


def choice(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.space_after = Pt(6)
    style_run(p.add_run("[  ] "), size=12, color=BLUE, bold=True)
    style_run(p.add_run(text), size=11.5)


def timeline(doc):
    rows = [
        ("5 min", "Começar", "Marcar ferramentas que você conhece."),
        ("8 min", "Aprender", "Ler dois exemplos: e-mail e chat."),
        ("10 min", "Escolher", "Decidir o canal para quatro situações."),
        ("15 min", "Montar", "Completar uma mensagem com frases-modelo."),
        ("8 min", "Revisar", "Usar o checklist com uma dupla."),
        ("6 min", "Finalizar", "Escrever ou copiar a mensagem final."),
        ("3 min", "Fechar", "Fazer uma autoavaliação simples."),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = (0.85, 1.25, 4.4)
    for idx, text in enumerate(("Tempo", "Etapa", "O que fazer")):
        cell = table.rows[0].cells[idx]
        cell.width = Inches(widths[idx])
        set_cell_shading(cell, BLUE)
        set_cell_margins(cell, 100, 100, 100, 100)
        style_run(cell.paragraphs[0].add_run(text), size=10, color="FFFFFF", bold=True)
    for tempo, etapa, acao in rows:
        cells = table.add_row().cells
        for idx, text in enumerate((tempo, etapa, acao)):
            cells[idx].width = Inches(widths[idx])
            set_cell_margins(cells[idx], 100, 100, 100, 100)
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            style_run(cells[idx].paragraphs[0].add_run(text), size=10)
    doc.add_paragraph()


def build():
    doc = Document()
    configure_document(doc)
    header = doc.sections[0].header.paragraphs[0]
    header.clear()
    style_run(header.add_run("SENAI | Atividade fácil - Ferramentas Digitais"), size=8.5, color=GRAY, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(45)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_run(p.add_run("ATIVIDADE FÁCIL"), size=12, color=ORANGE, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(6)
    style_run(p.add_run("Escolha o canal e monte uma mensagem"), size=24, color=BLUE, bold=True, font="Aptos Display")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    style_run(p.add_run("Ferramentas Digitais para Comunicação | 55 minutos + 5 minutos de apoio"), size=11.5, color=GRAY, italic=True)

    table = doc.add_table(rows=2, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row_idx, labels in enumerate((("Nome:", "Turma:"), ("Dupla de apoio:", "Data:"))):
        for col_idx, label in enumerate(labels):
            cell = table.rows[row_idx].cells[col_idx]
            cell.width = Inches(3.25)
            set_cell_margins(cell, 100, 120, 300, 120)
            style_run(cell.paragraphs[0].add_run(label), size=10, color=GRAY, bold=True)
    doc.add_paragraph()

    add_callout(doc, "Sua missão", "Escolher entre E-MAIL e CHAT. Depois, completar uma mensagem curta, clara e educada. Você pode ler em silêncio, apontar, marcar ou pedir ajuda à sua dupla.", fill=LIGHT_BLUE)
    doc.add_heading("O que você vai aprender", level=2)
    add_bullets(doc, [
        "Quando usar e-mail.",
        "Quando usar chat.",
        "Como escrever uma mensagem com quatro partes.",
        "Como revisar antes de enviar.",
    ])
    doc.add_heading("Tempo da atividade", level=2)
    timeline(doc)
    doc.add_page_break()

    doc.add_heading("PASSO 1 - O que eu já conheço?", level=1)
    add_body(doc, "Marque as ferramentas que você conhece. Não existe resposta errada.")
    for text in ("E-mail", "Chat ou mensagem", "Reunião por vídeo", "Documento compartilhado", "Quadro de tarefas"):
        choice(doc, text)
    add_body(doc, "Se quiser, escreva o nome de uma ferramenta que você usa:")
    line(doc)

    doc.add_heading("PASSO 2 - E-mail ou chat?", level=1)
    add_two_column_table(doc, ("E-MAIL", "CHAT"), [
        ("Usado para mensagem mais formal.", "Usado para mensagem curta e rápida."),
        ("Ajuda a guardar o registro.", "Ajuda a tirar uma dúvida simples."),
        ("Pode ter assunto e anexo.", "Precisa de contexto mesmo sendo curto."),
        ("Exemplo: enviar um trabalho.", "Exemplo: pedir o link de uma pasta."),
    ])
    add_callout(doc, "Dica simples", "Se a mensagem é importante e precisa ficar registrada, escolha E-MAIL. Se é uma dúvida curta durante a tarefa, escolha CHAT.", fill="FFF4E5", accent=ORANGE)

    doc.add_heading("PASSO 3 - Escolha o canal", level=1)
    situations = [
        ("1. Enviar um trabalho ao professor.", "E-MAIL", "CHAT"),
        ("2. Pedir rapidamente o link da pasta ao colega.", "E-MAIL", "CHAT"),
        ("3. Confirmar oficialmente a data de uma apresentação.", "E-MAIL", "CHAT"),
        ("4. Avisar ao grupo que você terminou uma tarefa.", "E-MAIL", "CHAT"),
    ]
    for situation, option_a, option_b in situations:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(7)
        style_run(p.add_run(situation), size=11.5, color=BLUE, bold=True)
        choice(doc, option_a)
        choice(doc, option_b)
        line(doc, "Por quê?")
    doc.add_page_break()

    doc.add_heading("PASSO 4 - Monte uma mensagem", level=1)
    add_callout(doc, "Situação", "Você não encontrou o link da pasta da atividade. Envie uma mensagem curta para a colega Ana pedindo o link até as 14h.", fill=LIGHT_BLUE)
    doc.add_heading("Uma mensagem clara tem quatro partes", level=2)
    add_two_column_table(doc, ("Parte", "Frase-modelo"), [
        ("1. Cumprimento", "Olá, Ana."),
        ("2. Contexto", "Estou fazendo a atividade de ferramentas digitais."),
        ("3. Pedido", "Você pode me enviar o link da pasta?"),
        ("4. Prazo e agradecimento", "Preciso dele até as 14h. Obrigado!"),
    ], widths=(1.8, 4.7))

    doc.add_heading("Escolha uma frase em cada parte", level=2)
    add_body(doc, "1. Cumprimento")
    choice(doc, "Olá, Ana.")
    choice(doc, "Bom dia, Ana.")
    add_body(doc, "2. Contexto")
    choice(doc, "Estou fazendo a atividade de ferramentas digitais.")
    choice(doc, "Estou organizando os arquivos do nosso grupo.")
    add_body(doc, "3. Pedido")
    choice(doc, "Você pode me enviar o link da pasta?")
    choice(doc, "Poderia compartilhar o link comigo?")
    add_body(doc, "4. Prazo e agradecimento")
    choice(doc, "Preciso dele até as 14h. Obrigado!")
    choice(doc, "Quando puder, envie até as 14h. Obrigado pela ajuda!")

    doc.add_heading("Copie as quatro frases escolhidas", level=2)
    for _ in range(7):
        line(doc)

    doc.add_heading("Banco de palavras", level=2)
    add_callout(doc, "Você pode usar", "olá | bom dia | atividade | pasta | link | enviar | compartilhar | por favor | até as 14h | obrigado", fill="F3F5F7")
    doc.add_page_break()

    doc.add_heading("PASSO 5 - Revise com uma dupla", level=1)
    add_body(doc, "A dupla pode ler, apontar ou apenas marcar. Não é obrigatório ler em voz alta.")
    checks = [
        "A mensagem tem um cumprimento.",
        "A mensagem explica o motivo.",
        "O pedido está claro.",
        "O prazo aparece.",
        "A mensagem usa palavras educadas.",
        "Não existem senha ou dados pessoais.",
    ]
    for item in checks:
        choice(doc, item)
    add_body(doc, "Uma ajuda que minha dupla deu:")
    line(doc)
    line(doc)

    doc.add_heading("PASSO 6 - Minha mensagem final", level=1)
    add_body(doc, "Escreva ou copie a versão final. Se precisar, use o modelo do Passo 4.")
    for _ in range(9):
        line(doc)

    doc.add_heading("PASSO 7 - Como foi para mim?", level=1)
    add_body(doc, "Marque uma opção em cada linha.")
    add_body(doc, "Escolher entre e-mail e chat foi:")
    choice(doc, "Fácil")
    choice(doc, "Mais ou menos")
    choice(doc, "Preciso de mais ajuda")
    add_body(doc, "Montar a mensagem foi:")
    choice(doc, "Fácil")
    choice(doc, "Mais ou menos")
    choice(doc, "Preciso de mais ajuda")
    add_body(doc, "Hoje eu consegui:")
    choice(doc, "Fazer sozinho")
    choice(doc, "Fazer com ajuda")
    choice(doc, "Começar e tentar")

    doc.add_heading("Avaliação do professor", level=1)
    add_two_column_table(doc, ("Critério", "Registro"), [
        ("Escolheu canais adequados", "[  ] Sim   [  ] Com apoio   [  ] Ainda não"),
        ("Montou as quatro partes", "[  ] Sim   [  ] Com apoio   [  ] Ainda não"),
        ("Usou linguagem respeitosa", "[  ] Sim   [  ] Com apoio   [  ] Ainda não"),
        ("Revisou antes de finalizar", "[  ] Sim   [  ] Com apoio   [  ] Ainda não"),
    ], widths=(3.25, 3.25))
    add_callout(doc, "Importante", "O sucesso nesta atividade é participar, fazer escolhas e construir uma mensagem compreensível, com ou sem apoio.", fill=LIGHT_BLUE)

    props = doc.core_properties
    props.title = "Atividade Fácil - Ferramentas Digitais para Comunicação"
    props.subject = "Atividade inclusiva de até 1 hora para estudantes de 14 anos"
    props.author = "SENAI - Programa Rio do Sul Mais Tech"
    props.keywords = "atividade fácil, comunicação digital, inclusão, e-mail, chat"

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
