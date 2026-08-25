from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
TARGET = ROOT / "sistema" / "FICHA-PRODUTO-MAIS-TECH" / "INTRODUCAO_COMUNICACAO_ORAL_ESCRITA" / "04 - FERRAMENTAS-DIGITAIS" / "AtividadesFerramentasDigitaisFacil.docx"


def remove_paragraph(paragraph):
    element = paragraph._element
    element.getparent().remove(element)


def remove_manual_page_breaks(paragraph):
    for br in list(paragraph._p.xpath(".//w:br[@w:type='page']")):
        br.getparent().remove(br)
    for rendered_break in list(paragraph._p.xpath(".//w:lastRenderedPageBreak")):
        rendered_break.getparent().remove(rendered_break)


def compact_paragraph(paragraph, in_table=False):
    remove_manual_page_breaks(paragraph)
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(1 if not in_table else 0)
    fmt.line_spacing = 1.0
    fmt.keep_together = False

    style_name = paragraph.style.name if paragraph.style else ""
    if style_name == "Heading 1":
        fmt.space_before = Pt(5)
        fmt.space_after = Pt(2)
        fmt.keep_with_next = True
        target_size = 13
    elif style_name == "Heading 2":
        fmt.space_before = Pt(3)
        fmt.space_after = Pt(1)
        fmt.keep_with_next = True
        target_size = 10.5
    elif style_name == "Heading 3":
        fmt.space_before = Pt(2)
        fmt.space_after = Pt(1)
        fmt.keep_with_next = True
        target_size = 9.5
    else:
        target_size = 8.8 if not in_table else 8.2

    for run in paragraph.runs:
        current = run.font.size.pt if run.font.size else None
        if current is None or current > target_size:
            run.font.size = Pt(target_size)


def compact_table(table):
    table.autofit = True
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "5000")
    tbl_w.set(qn("w:type"), "pct")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "autofit")

    for row in table.rows:
        row.height = None
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_mar = tc_pr.find(qn("w:tcMar"))
            if tc_mar is None:
                tc_mar = OxmlElement("w:tcMar")
                tc_pr.append(tc_mar)
            for margin in ("top", "start", "bottom", "end"):
                node = tc_mar.find(qn(f"w:{margin}"))
                if node is None:
                    node = OxmlElement(f"w:{margin}")
                    tc_mar.append(node)
                node.set(qn("w:w"), "40")
                node.set(qn("w:type"), "dxa")
            for paragraph in cell.paragraphs:
                compact_paragraph(paragraph, in_table=True)


def compact_document():
    if not TARGET.exists():
        raise FileNotFoundError(TARGET)

    doc = Document(TARGET)
    for section in doc.sections:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Inches(11)
        section.page_height = Inches(8.5)
        section.top_margin = Inches(0.3)
        section.bottom_margin = Inches(0.3)
        section.left_margin = Inches(0.3)
        section.right_margin = Inches(0.3)
        section.header_distance = Inches(0.15)
        section.footer_distance = Inches(0.15)

        for paragraph in section.header.paragraphs:
            compact_paragraph(paragraph, in_table=True)
            for run in paragraph.runs:
                run.font.size = Pt(7)
        for paragraph in section.footer.paragraphs:
            compact_paragraph(paragraph, in_table=True)
            for run in paragraph.runs:
                run.font.size = Pt(7)

    normal = doc.styles["Normal"]
    normal.font.size = Pt(8.8)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(1)
    normal.paragraph_format.line_spacing = 1.0

    for style_name, size in (("Heading 1", 13), ("Heading 2", 10.5), ("Heading 3", 9.5)):
        style = doc.styles[style_name]
        style.font.size = Pt(size)
        style.paragraph_format.space_before = Pt(3)
        style.paragraph_format.space_after = Pt(1)
        style.paragraph_format.line_spacing = 1.0

    for paragraph in list(doc.paragraphs):
        if not paragraph.text.strip() and not paragraph._p.xpath(".//w:drawing|.//w:pict"):
            remove_paragraph(paragraph)
        else:
            compact_paragraph(paragraph)

    for table in doc.tables:
        compact_table(table)

    doc.save(TARGET)
    print(TARGET)


if __name__ == "__main__":
    compact_document()
