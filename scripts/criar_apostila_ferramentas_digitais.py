from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "sistema" / "FICHA-PRODUTO-MAIS-TECH" / "INTRODUCAO_COMUNICACAO_ORAL_ESCRITA" / "Ferramentas Digitais para Comunicação.docx"

BLUE = "004384"
LIGHT_BLUE = "E8F0FE"
ORANGE = "F7941D"
DARK = "202124"
GRAY = "666666"
LIGHT_GRAY = "F3F5F7"
GREEN = "2E7D32"
RED = "C62828"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_repeat_keep(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    keep_next = OxmlElement("w:keepNext")
    p_pr.append(keep_next)


def add_page_field(paragraph):
    paragraph.add_run("Página ")
    run = paragraph.add_run()
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char_1, instr_text, fld_char_2])


def style_run(run, size=11, color=DARK, bold=False, italic=False, font="Aptos"):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def add_body(doc, text, bold_lead=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.25
    if bold_lead and text.startswith(bold_lead):
        style_run(p.add_run(bold_lead), bold=True)
        style_run(p.add_run(text[len(bold_lead):]))
    else:
        style_run(p.add_run(text))
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.188)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        style_run(p.add_run(item))


def add_steps(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.188)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        style_run(p.add_run(item))


def add_callout(doc, title, text, fill=LIGHT_BLUE, accent=BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(6.5)
    cell = table.cell(0, 0)
    cell.width = Inches(6.5)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_shading(cell, fill)
    set_cell_margins(cell, 130, 180, 130, 180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    style_run(p.add_run(title), size=11, color=accent, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.18
    style_run(p2.add_run(text), size=10.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_activity(doc, title, objective, steps, deliverable):
    doc.add_heading(title, level=2)
    add_callout(doc, "Objetivo da atividade", objective, fill="FFF4E5", accent=ORANGE)
    add_steps(doc, steps)
    add_body(doc, f"Entrega: {deliverable}", bold_lead="Entrega:")


def add_two_column_table(doc, headers, rows, widths=(3.25, 3.25)):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.width = Inches(widths[idx])
        set_cell_shading(cell, BLUE)
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        style_run(p.add_run(header), size=10, color="FFFFFF", bold=True)
    set_repeat_table_header(table.rows[0])
    for left, right in rows:
        cells = table.add_row().cells
        for idx, text in enumerate((left, right)):
            cells[idx].width = Inches(widths[idx])
            set_cell_margins(cells[idx])
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            style_run(p.add_run(text), size=10)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.82)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(DARK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 17, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 11.5, "1F4D78", 9, 4),
    ):
        style = styles[name]
        style.font.name = "Aptos Display"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Aptos Display")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos Display")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    style_run(hp.add_run("SENAI | Ferramentas Digitais para Comunicação"), size=8.5, color=GRAY, bold=True)
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    style_run(fp.add_run("Rio do Sul Mais Tech  |  "), size=8.5, color=GRAY)
    add_page_field(fp)


def cover(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(90)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_run(p.add_run("APOSTILA"), size=12, color=ORANGE, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(10)
    style_run(p.add_run("Ferramentas Digitais\npara Comunicação"), size=29, color=BLUE, bold=True, font="Aptos Display")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(28)
    style_run(p.add_run("Comunicação profissional, colaboração, netiqueta e segurança digital"), size=13, color=GRAY, italic=True)
    add_callout(doc, "Unidade Curricular", "Introdução à Comunicação Oral e Escrita para o Mundo do Trabalho | Carga horária da UC: 33 horas", fill=LIGHT_BLUE)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(65)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_run(p.add_run("Programa Rio do Sul Mais Tech"), size=12, color=BLUE, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_run(p.add_run("SENAI / Prefeitura Municipal de Rio do Sul\nPúblico-alvo: estudantes do 8º e 9º ano"), size=10.5, color=GRAY)
    doc.add_page_break()


def introduction(doc):
    doc.add_heading("Apresentação", level=1)
    add_body(doc, "No mundo do trabalho, comunicar-se bem não depende apenas de escrever corretamente ou falar com clareza. Também é necessário escolher a ferramenta digital adequada, organizar informações, proteger dados e respeitar as pessoas em cada interação.")
    add_body(doc, "Esta apostila apresenta situações práticas de comunicação digital. O objetivo não é decorar botões de aplicativos, mas desenvolver critérios para decidir qual canal usar, como produzir uma mensagem profissional e como agir com segurança e responsabilidade.")
    add_callout(doc, "Competência central", "Utilizar ferramentas digitais de comunicação corporativa com responsabilidade, clareza, organização e segurança.", fill="FFF4E5", accent=ORANGE)
    doc.add_heading("Como utilizar esta apostila", level=2)
    add_bullets(doc, [
        "Leia os exemplos e compare as escolhas de linguagem e canal.",
        "Realize as atividades usando contas e ambientes autorizados pelo professor.",
        "Nunca compartilhe senhas, dados pessoais ou arquivos reais sem permissão.",
        "Guarde as produções em uma pasta digital para formar seu portfólio.",
        "Ao final de cada unidade, use o checklist para revisar sua aprendizagem.",
    ])
    doc.add_heading("Mapa de aprendizagem", level=2)
    add_two_column_table(doc, ("Unidade", "O que você aprenderá"), [
        ("1. Escolha do canal", "Relacionar objetivo, urgência, público e registro necessário."),
        ("2. E-mail profissional", "Planejar, escrever, revisar e responder mensagens corporativas."),
        ("3. Colaboração em nuvem", "Compartilhar arquivos, editar em equipe e controlar versões."),
        ("4. Mensagens e reuniões", "Usar chats e videoconferências com objetividade e respeito."),
        ("5. Organização do trabalho", "Comunicar tarefas, prazos, responsáveis e progresso."),
        ("6. Netiqueta e segurança", "Proteger contas, reconhecer golpes e cuidar da privacidade."),
        ("7. Projeto integrador", "Planejar um fluxo completo de comunicação para uma equipe."),
    ], widths=(1.85, 4.65))
    doc.add_page_break()


def unit_channel(doc):
    doc.add_heading("1. Escolhendo a ferramenta certa", level=1)
    add_body(doc, "Uma ferramenta digital é um meio para alcançar um objetivo. Antes de abrir um aplicativo, pergunte: quem precisa receber a informação, qual é a urgência, a mensagem precisa ficar registrada e haverá necessidade de editar um arquivo em conjunto?")
    add_two_column_table(doc, ("Situação", "Canal mais adequado"), [
        ("Aviso formal que precisa ficar registrado", "E-mail corporativo ou comunicado em plataforma oficial."),
        ("Dúvida rápida durante uma tarefa", "Chat da equipe, com contexto e pergunta objetiva."),
        ("Decisão complexa com várias pessoas", "Reunião por vídeo, seguida de resumo escrito."),
        ("Produção conjunta de um documento", "Editor colaborativo em nuvem."),
        ("Acompanhamento de prazos e responsáveis", "Quadro ou aplicativo de gestão de tarefas."),
        ("Emergência operacional", "Canal imediato definido pela organização, seguido de registro."),
    ])
    doc.add_heading("Critérios de decisão", level=2)
    add_bullets(doc, [
        "Urgência: a pessoa precisa agir agora, hoje ou nesta semana?",
        "Formalidade: a mensagem representa uma decisão, solicitação ou compromisso?",
        "Complexidade: o assunto pode ser entendido em poucas linhas?",
        "Rastreabilidade: será necessário consultar a conversa posteriormente?",
        "Público: todos têm acesso ao canal e sabem utilizá-lo?",
        "Segurança: o canal é autorizado para o tipo de dado compartilhado?",
    ])
    add_callout(doc, "Regra prática", "Quanto mais importante for uma decisão, maior deve ser a preocupação com registro, clareza, confirmação e segurança.")
    add_activity(doc, "Atividade 1 — Qual canal usar?", "Justificar a escolha da ferramenta em situações profissionais.", [
        "Forme uma dupla e leia seis situações propostas pelo professor.",
        "Escolha um canal para cada situação.",
        "Registre a justificativa considerando urgência, formalidade e necessidade de registro.",
        "Compare as respostas com outra dupla e revise uma escolha.",
    ], "Tabela com situação, canal escolhido e justificativa.")
    doc.add_page_break()


def unit_email(doc):
    doc.add_heading("2. E-mail profissional", level=1)
    add_body(doc, "O e-mail profissional é usado para solicitações, confirmações, envio de documentos e comunicações que precisam permanecer registradas. Uma boa mensagem permite que o destinatário entenda rapidamente o assunto e saiba qual ação realizar.")
    doc.add_heading("Estrutura essencial", level=2)
    add_steps(doc, [
        "Assunto: resuma o objetivo com palavras específicas.",
        "Saudação: cumprimente de acordo com o contexto e a relação profissional.",
        "Contexto: explique em uma ou duas frases por que está escrevendo.",
        "Ação esperada: diga claramente o que precisa ser feito e até quando.",
        "Encerramento: agradeça e indique disponibilidade para esclarecimentos.",
        "Assinatura: informe nome, turma, equipe ou função quando necessário.",
    ])
    add_two_column_table(doc, ("Evite", "Prefira"), [
        ("Assunto: Oi", "Assunto: Confirmação da apresentação de 18/08"),
        ("Me manda aquilo.", "Poderia enviar a versão revisada do relatório até as 15h?"),
        ("URGENTE!!!", "Prazo hoje: aprovação do comunicado"),
        ("Texto inteiro em letras maiúsculas", "Frases curtas, pontuação adequada e destaques moderados"),
        ("Anexo sem explicação", "Identificação do arquivo e breve descrição do conteúdo"),
    ])
    doc.add_heading("Para, CC e CCO", level=2)
    add_bullets(doc, [
        "Para: pessoas responsáveis por responder ou agir.",
        "CC: pessoas que precisam acompanhar a informação.",
        "CCO: oculta endereços dos demais destinatários; use quando a privacidade exigir.",
        "Responder a todos: use apenas quando sua resposta realmente interessar a todo o grupo.",
    ])
    add_callout(doc, "Antes de enviar", "Confira destinatários, assunto, tom, prazo, anexos e dados sensíveis. Muitos erros profissionais acontecem por pressa, não por falta de conhecimento.", fill="FFF4E5", accent=ORANGE)
    doc.add_heading("Modelo de e-mail", level=2)
    add_callout(doc, "Assunto: Envio do roteiro para apresentação de sexta-feira", "Olá, professora Ana. Conforme combinado, envio em anexo o roteiro da apresentação sobre segurança digital. Poderia confirmar o recebimento e informar se precisamos realizar algum ajuste até quinta-feira? Obrigado. Atenciosamente, Lucas - Turma 9B.")
    add_activity(doc, "Atividade 2 — Escrevendo um e-mail profissional", "Produzir uma solicitação clara, respeitosa e completa.", [
        "Escolha uma situação: solicitar informação, enviar trabalho ou confirmar reunião.",
        "Escreva assunto, saudação, contexto, ação esperada, encerramento e assinatura.",
        "Troque o texto com um colega e aplique o checklist antes do envio simulado.",
        "Reescreva pelo menos um trecho para deixá-lo mais objetivo.",
    ], "E-mail revisado em documento ou ambiente de simulação indicado pelo professor.")
    doc.add_page_break()


def unit_cloud(doc):
    doc.add_heading("3. Colaboração em nuvem", level=1)
    add_body(doc, "Google Workspace e Microsoft 365 reúnem editores de texto, planilhas, apresentações, armazenamento, calendário e reuniões. Essas plataformas permitem que uma equipe trabalhe no mesmo arquivo sem criar diversas cópias desconectadas.")
    add_two_column_table(doc, ("Google Workspace", "Microsoft 365"), [
        ("Drive: arquivos e pastas", "OneDrive e SharePoint: arquivos e sites de equipe"),
        ("Docs: documentos colaborativos", "Word para Web: documentos colaborativos"),
        ("Sheets: planilhas", "Excel para Web: planilhas"),
        ("Slides: apresentações", "PowerPoint para Web: apresentações"),
        ("Meet e Chat: reuniões e mensagens", "Teams: reuniões, canais e mensagens"),
        ("Calendar: agenda compartilhada", "Outlook: e-mail e calendário"),
    ])
    doc.add_heading("Compartilhamento responsável", level=2)
    add_bullets(doc, [
        "Visualizador: pode consultar, mas não alterar.",
        "Comentador: pode sugerir e comentar sem editar diretamente.",
        "Editor: pode modificar o conteúdo; conceda apenas quando necessário.",
        "Link restrito: somente pessoas autorizadas acessam.",
        "Link público: pode expor informações; evite sem autorização.",
    ])
    doc.add_heading("Organização de arquivos", level=2)
    add_two_column_table(doc, ("Prática", "Exemplo"), [
        ("Nome claro", "Relatorio_visita_tecnica_2026-08-20.docx"),
        ("Pasta por projeto", "Projeto Feira / Pesquisa / Imagens / Entregas"),
        ("Versão controlada", "Usar histórico da plataforma em vez de Final_final_2"),
        ("Responsável definido", "Uma pessoa revisa permissões e organiza a pasta"),
        ("Arquivamento", "Mover entregas concluídas para pasta identificada"),
    ])
    add_callout(doc, "Comentário x edição", "Use comentários para perguntar, justificar ou sugerir. Edite diretamente quando tiver autorização e quando a mudança não apagar a contribuição de outra pessoa.")
    add_activity(doc, "Atividade 3 — Documento colaborativo", "Praticar compartilhamento, comentários e revisão em equipe.", [
        "Crie um documento sobre boas práticas de comunicação da turma.",
        "Compartilhe com dois colegas usando permissões diferentes.",
        "Cada colega deve inserir um comentário e uma sugestão.",
        "Consulte o histórico de versões e identifique uma alteração.",
        "Revise as permissões e remova acessos desnecessários.",
    ], "Documento colaborativo organizado e registro das permissões utilizadas.")
    doc.add_page_break()


def unit_messaging_meetings(doc):
    doc.add_heading("4. Chats, canais e videoconferências", level=1)
    add_body(doc, "Mensagens instantâneas são úteis para comunicações breves, mas a rapidez não elimina a necessidade de contexto e respeito. Em canais profissionais, uma mensagem deve permitir que outra pessoa compreenda o assunto mesmo que não tenha acompanhado a conversa desde o início.")
    doc.add_heading("Mensagem objetiva em quatro partes", level=2)
    add_steps(doc, [
        "Cumprimento ou identificação da pessoa/equipe.",
        "Contexto breve: projeto, tarefa ou problema.",
        "Pergunta ou ação esperada.",
        "Prazo ou nível de urgência, quando existir.",
    ])
    add_two_column_table(doc, ("Mensagem confusa", "Mensagem profissional"), [
        ("Oi. Deu errado.", "Bom dia, equipe. O link do formulário da pesquisa retorna acesso negado. Algué com permissão pode revisar o compartilhamento até as 14h?"),
        ("Você viu?", "Marina, você conseguiu revisar o slide 5 da apresentação? Precisamos fechar a versão hoje."),
        ("Me responde urgente", "Preciso confirmar o horário da reunião antes das 11h para avisar a turma."),
    ])
    doc.add_heading("Boas práticas em videoconferências", level=2)
    add_bullets(doc, [
        "Teste microfone, câmera e conexão antes do horário.",
        "Entre com nome identificável e mantenha o microfone fechado quando não estiver falando.",
        "Use chat, levantar a mão e reações sem interromper quem apresenta.",
        "Evite expor informações pessoais no fundo da imagem ou no compartilhamento de tela.",
        "Peça autorização antes de gravar.",
        "Finalize com decisões, responsáveis e prazos registrados.",
    ])
    add_activity(doc, "Atividade 4 — Reunião digital de 10 minutos", "Organizar uma reunião curta com pauta, participação e registro.", [
        "Defina objetivo, três itens de pauta e tempo de cada item.",
        "Distribua papéis: facilitador, controlador de tempo e responsável pelo registro.",
        "Realize a reunião usando um ambiente autorizado.",
        "Produza um resumo com decisões, responsáveis e prazos.",
    ], "Pauta e resumo da reunião em um único documento.")
    doc.add_page_break()


def unit_tasks(doc):
    doc.add_heading("5. Comunicação e gestão de tarefas", level=1)
    add_body(doc, "Ferramentas como Trello, Microsoft Planner, Google Tasks, Asana e quadros semelhantes tornam o trabalho visível. Elas não servem apenas para listar tarefas: comunicam o que deve ser feito, por quem, até quando e em qual etapa o trabalho se encontra.")
    doc.add_heading("Anatomia de uma boa tarefa", level=2)
    add_bullets(doc, [
        "Título com verbo de ação: Revisar roteiro da apresentação.",
        "Descrição com resultado esperado e critérios de conclusão.",
        "Responsável principal, evitando tarefas sem dono.",
        "Prazo realista e prioridade definida.",
        "Anexos, links e contexto necessários.",
        "Status atualizado: a fazer, em andamento, bloqueada ou concluída.",
    ])
    add_two_column_table(doc, ("Pouco claro", "Claro e verificável"), [
        ("Fazer slides", "Criar cinco slides sobre phishing usando o modelo da turma até quarta-feira."),
        ("Pesquisar", "Selecionar três fontes confiáveis sobre privacidade e registrar os links."),
        ("Ver com o grupo", "Revisar o roteiro com Ana e Pedro e registrar ajustes no documento."),
        ("Terminar logo", "Enviar a versão revisada até 16h de sexta-feira."),
    ])
    add_callout(doc, "Status não é julgamento", "Marcar uma tarefa como bloqueada não significa fracasso. Significa comunicar cedo que existe um impedimento e permitir que a equipe ajude.", fill="FFF4E5", accent=ORANGE)
    add_activity(doc, "Atividade 5 — Quadro de projeto", "Transformar um objetivo em tarefas comunicáveis e acompanháveis.", [
        "Escolha um pequeno projeto da turma.",
        "Crie colunas A fazer, Em andamento, Em revisão e Concluído.",
        "Cadastre pelo menos seis tarefas com responsável, prazo e descrição.",
        "Simule uma tarefa bloqueada e escreva uma solicitação de ajuda objetiva.",
        "Atualize o quadro e apresente o progresso em dois minutos.",
    ], "Quadro digital ou representação autorizada com tarefas completas.")
    doc.add_page_break()


def unit_netiquette(doc):
    doc.add_heading("6. Netiqueta e cidadania digital", level=1)
    add_body(doc, "Netiqueta é o conjunto de atitudes que favorece convivência respeitosa e eficiente em ambientes digitais. A tela não elimina o impacto das palavras. Mensagens podem ser encaminhadas, registradas e interpretadas sem o apoio do tom de voz ou da expressão facial.")
    add_bullets(doc, [
        "Escreva como se a mensagem pudesse ser lida por toda a equipe.",
        "Discorde de ideias sem atacar pessoas.",
        "Evite ironia, excesso de emojis e abreviações em contextos formais.",
        "Respeite horários e não cobre resposta imediata sem necessidade real.",
        "Não encaminhe conversas, imagens ou arquivos sem autorização.",
        "Corrija informações erradas de forma transparente.",
        "Dê crédito a autores, imagens e fontes utilizadas.",
    ])
    doc.add_heading("Conflitos digitais", level=2)
    add_steps(doc, [
        "Pare antes de responder quando estiver irritado.",
        "Releia a mensagem e diferencie fato, interpretação e sentimento.",
        "Peça esclarecimento em vez de presumir intenção negativa.",
        "Mude para conversa por voz ou reunião se o assunto for complexo.",
        "Registre acordos importantes por escrito após a conversa.",
        "Procure um responsável quando houver ofensa, ameaça ou exposição indevida.",
    ])
    add_activity(doc, "Atividade 6 — Reescrita com netiqueta", "Reconhecer riscos de tom e reescrever mensagens de forma assertiva.", [
        "Leia mensagens com pressão, ironia, acusação ou falta de contexto.",
        "Identifique o problema de cada mensagem.",
        "Reescreva usando fato, necessidade, pedido claro e prazo.",
        "Explique como a nova versão reduz conflitos.",
    ], "Quadro comparativo com mensagem original, risco e versão revisada.")
    doc.add_page_break()


def unit_security(doc):
    doc.add_heading("7. Segurança e privacidade na comunicação", level=1)
    add_body(doc, "Comunicação digital segura protege pessoas, projetos e organizações. Um clique em link falso ou o envio de um arquivo para o destinatário errado pode expor senhas, dados pessoais e informações de trabalho.")
    doc.add_heading("Proteção de contas", level=2)
    add_bullets(doc, [
        "Use senhas longas, únicas e difíceis de adivinhar.",
        "Ative autenticação em dois fatores sempre que possível.",
        "Não compartilhe códigos de verificação.",
        "Bloqueie a tela ao se afastar do dispositivo.",
        "Saia de contas abertas em computadores compartilhados.",
        "Mantenha aplicativos e sistemas atualizados.",
    ])
    doc.add_heading("Como reconhecer phishing", level=2)
    add_two_column_table(doc, ("Sinal de alerta", "Ação segura"), [
        ("Pressão para agir imediatamente", "Pare e confirme por outro canal oficial."),
        ("Remetente ou domínio estranho", "Compare cuidadosamente o endereço real."),
        ("Link encurtado ou diferente do texto", "Não clique; verifique o destino e acesse o site oficial."),
        ("Pedido de senha ou código", "Não informe; serviços legítimos não pedem senha por mensagem."),
        ("Anexo inesperado", "Confirme com o remetente antes de abrir."),
        ("Oferta boa demais", "Desconfie e pesquise em fontes oficiais."),
    ])
    doc.add_heading("Privacidade e dados pessoais", level=2)
    add_body(doc, "Nome completo, documento, endereço, telefone, localização, imagem e informações de saúde podem identificar ou expor uma pessoa. Compartilhe apenas o necessário, com finalidade clara, permissão e canal autorizado.")
    add_callout(doc, "Se você clicou em algo suspeito", "Desconecte-se quando necessário, avise imediatamente um professor ou responsável, altere a senha pelo site oficial e siga as orientações da instituição. Esconder o erro aumenta o risco.", fill="FDECEA", accent=RED)
    add_activity(doc, "Atividade 7 — Investigadores de phishing", "Analisar criticamente sinais de fraude sem acessar conteúdo perigoso.", [
        "Analise capturas simuladas fornecidas pelo professor.",
        "Marque remetente, urgência, link, pedido e linguagem suspeita.",
        "Classifique cada exemplo como confiável, suspeito ou fraude provável.",
        "Escreva a ação segura recomendada.",
    ], "Ficha de análise de mensagens simuladas.")
    doc.add_page_break()


def unit_accessibility(doc):
    doc.add_heading("8. Comunicação digital acessível", level=1)
    add_body(doc, "Acessibilidade digital significa produzir mensagens e documentos que possam ser compreendidos e utilizados por pessoas com diferentes necessidades. Clareza, organização e alternativas visuais beneficiam toda a equipe.")
    add_bullets(doc, [
        "Use títulos e subtítulos reais para organizar documentos.",
        "Escreva textos alternativos que expliquem imagens importantes.",
        "Não dependa somente de cor para transmitir informação.",
        "Mantenha bom contraste entre texto e fundo.",
        "Use links descritivos, como Acessar o roteiro da atividade.",
        "Adicione legendas ou transcrição a conteúdos em vídeo e áudio.",
        "Prefira frases diretas, parágrafos curtos e listas bem organizadas.",
    ])
    add_activity(doc, "Atividade 8 — Revisão de acessibilidade", "Aplicar um checklist de acessibilidade a uma comunicação digital.", [
        "Escolha um documento ou apresentação produzida anteriormente.",
        "Revise hierarquia de títulos, contraste, links e imagens.",
        "Implemente pelo menos quatro melhorias.",
        "Registre antes, depois e justificativa de cada ajuste.",
    ], "Registro das melhorias de acessibilidade realizadas.")
    doc.add_page_break()


def project_and_assessment(doc):
    doc.add_heading("9. Projeto integrador", level=1)
    add_body(doc, "Em equipe, planejem a comunicação digital de um evento escolar, uma feira de profissões ou uma campanha de conscientização. O projeto deve demonstrar escolha de canais, produção de mensagens, colaboração, organização de tarefas e segurança.")
    doc.add_heading("Entregas obrigatórias", level=2)
    add_steps(doc, [
        "Mapa de públicos e canais: quem precisa receber qual informação.",
        "E-mail profissional de convite ou solicitação.",
        "Pasta compartilhada com estrutura, nomes e permissões adequadas.",
        "Documento colaborativo com comentários e histórico de contribuições.",
        "Quadro de tarefas com responsáveis, prazos e status.",
        "Pauta e resumo de uma reunião digital.",
        "Checklist de netiqueta, segurança, privacidade e acessibilidade.",
        "Apresentação final de cinco minutos com demonstração do fluxo criado.",
    ])
    doc.add_heading("Critérios de avaliação", level=2)
    add_two_column_table(doc, ("Critério", "Evidência esperada"), [
        ("Clareza e adequação", "Mensagens objetivas, respeitosas e apropriadas ao público."),
        ("Escolha de ferramentas", "Canais justificados por urgência, registro e complexidade."),
        ("Organização", "Arquivos, tarefas, responsáveis e prazos facilmente localizáveis."),
        ("Colaboração", "Contribuições equilibradas, comentários e revisão em equipe."),
        ("Segurança e privacidade", "Permissões corretas e ausência de dados pessoais desnecessários."),
        ("Acessibilidade", "Estrutura clara, contraste, links e alternativas para conteúdos visuais."),
        ("Apresentação", "Explicação clara das decisões e aprendizados."),
    ], widths=(2.0, 4.5))
    doc.add_heading("Autoavaliação", level=2)
    add_bullets(doc, [
        "Consigo escolher entre e-mail, chat, reunião, documento e quadro de tarefas?",
        "Reviso destinatários, anexos, tom e dados sensíveis antes de enviar?",
        "Sei compartilhar arquivos com a permissão adequada?",
        "Comunico bloqueios, prazos e pedidos de ajuda com clareza?",
        "Reconheço sinais de phishing e sei a quem pedir ajuda?",
        "Produzo documentos que outras pessoas conseguem compreender e acessar?",
    ])
    doc.add_page_break()


def glossary(doc):
    doc.add_heading("Glossário essencial", level=1)
    terms = [
        ("Anexo", "Arquivo enviado junto a uma mensagem."),
        ("Autenticação em dois fatores", "Segunda verificação além da senha."),
        ("Canal", "Meio utilizado para transmitir uma mensagem."),
        ("CC / CCO", "Cópia visível e cópia oculta em e-mails."),
        ("Computação em nuvem", "Uso de arquivos e serviços acessados pela internet."),
        ("Controle de versão", "Registro das alterações feitas em um arquivo."),
        ("Netiqueta", "Boas práticas de convivência e comunicação digital."),
        ("Permissão", "Nível de acesso concedido a uma pessoa."),
        ("Phishing", "Tentativa de enganar para roubar dados ou instalar conteúdo malicioso."),
        ("Rastreabilidade", "Possibilidade de consultar o histórico de mensagens e decisões."),
        ("Sincrônico", "Interação que acontece ao mesmo tempo, como videoconferência."),
        ("Assíncrono", "Interação que permite respostas em momentos diferentes, como e-mail."),
    ]
    add_two_column_table(doc, ("Termo", "Significado"), terms, widths=(2.0, 4.5))
    doc.add_heading("Checklist final de comunicação digital", level=1)
    add_bullets(doc, [
        "Escolhi o canal adequado ao objetivo e à urgência.",
        "A mensagem apresenta contexto e ação esperada.",
        "O tom é respeitoso, profissional e adequado ao público.",
        "Destinatários, links, anexos e permissões foram conferidos.",
        "Dados pessoais e informações sensíveis estão protegidos.",
        "A comunicação é acessível e pode ser compreendida sem contexto oculto.",
        "Decisões, responsáveis e prazos importantes ficaram registrados.",
    ])
    doc.add_heading("Referências e recursos sugeridos", level=1)
    add_bullets(doc, [
        "Google Workspace Learning Center — guias de Gmail, Drive, Docs, Meet e Calendar.",
        "Microsoft Support — guias de Outlook, OneDrive, Word, Teams e Planner.",
        "CERT.br — Cartilha de Segurança para Internet.",
        "SaferNet Brasil — materiais de cidadania e segurança digital.",
        "Materiais e políticas de uso de tecnologia disponibilizados pelo SENAI e pelo professor.",
    ])
    add_callout(doc, "Mensagem final", "A ferramenta pode mudar, mas os princípios permanecem: clareza, respeito, organização, responsabilidade e segurança.", fill=LIGHT_BLUE)


def build():
    doc = Document()
    configure_document(doc)
    cover(doc)
    introduction(doc)
    unit_channel(doc)
    unit_email(doc)
    unit_cloud(doc)
    unit_messaging_meetings(doc)
    unit_tasks(doc)
    unit_netiquette(doc)
    unit_security(doc)
    unit_accessibility(doc)
    project_and_assessment(doc)
    glossary(doc)

    properties = doc.core_properties
    properties.title = "Ferramentas Digitais para Comunicação"
    properties.subject = "Apostila didática para comunicação digital no mundo do trabalho"
    properties.author = "SENAI - Programa Rio do Sul Mais Tech"
    properties.keywords = "comunicação digital, e-mail, colaboração, netiqueta, segurança"
    properties.comments = "Material didático baseado na ementa da unidade curricular."

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
