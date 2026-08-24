from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from criar_apostila_ferramentas_digitais import (
    BLUE,
    DARK,
    GRAY,
    GREEN,
    LIGHT_BLUE,
    ORANGE,
    RED,
    add_body,
    add_bullets,
    add_callout,
    add_page_field,
    add_steps,
    add_two_column_table,
    configure_document,
    set_cell_margins,
    set_cell_shading,
    style_run,
)


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "sistema" / "FICHA-PRODUTO-MAIS-TECH" / "INTRODUCAO_COMUNICACAO_ORAL_ESCRITA" / "AtividadesFerramentasDigitais.docx"


def response_lines(doc, count=4):
    for _ in range(count):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(5)
        run = p.add_run("________________________________________________________________________________")
        style_run(run, size=9, color="A0A0A0")


def add_timeline(doc, rows):
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = (0.85, 1.8, 3.85)
    for idx, text in enumerate(("Tempo", "Etapa", "Ação")):
        cell = table.rows[0].cells[idx]
        cell.width = Inches(widths[idx])
        set_cell_shading(cell, BLUE)
        set_cell_margins(cell)
        style_run(cell.paragraphs[0].add_run(text), size=10, color="FFFFFF", bold=True)
    for tempo, etapa, acao in rows:
        cells = table.add_row().cells
        for idx, text in enumerate((tempo, etapa, acao)):
            cells[idx].width = Inches(widths[idx])
            set_cell_margins(cells[idx])
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            style_run(cells[idx].paragraphs[0].add_run(text), size=9.5)
    doc.add_paragraph()


def add_rubric(doc, rows):
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = (2.1, 3.65, 0.75)
    for idx, text in enumerate(("Critério", "O que observar", "Pontos")):
        cell = table.rows[0].cells[idx]
        cell.width = Inches(widths[idx])
        set_cell_shading(cell, "E8EEF5")
        set_cell_margins(cell)
        style_run(cell.paragraphs[0].add_run(text), size=9.5, color=BLUE, bold=True)
    for criterio, evidencia, pontos in rows:
        cells = table.add_row().cells
        for idx, text in enumerate((criterio, evidencia, pontos)):
            cells[idx].width = Inches(widths[idx])
            set_cell_margins(cells[idx])
            style_run(cells[idx].paragraphs[0].add_run(text), size=9.2)
    doc.add_paragraph()


def activity_header(doc, number, title, subtitle):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(3)
    style_run(p.add_run(f"ATIVIDADE {number}  |  60 MINUTOS"), size=10.5, color=ORANGE, bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    style_run(p.add_run(title), size=22, color=BLUE, bold=True, font="Aptos Display")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    style_run(p.add_run(subtitle), size=11.5, color=GRAY, italic=True)


def student_identification(doc):
    table = doc.add_table(rows=2, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    labels = (("Nome:", "Turma:"), ("Equipe:", "Data:"))
    for row_idx, row in enumerate(labels):
        for col_idx, label in enumerate(row):
            cell = table.rows[row_idx].cells[col_idx]
            cell.width = Inches(3.25)
            set_cell_margins(cell, 100, 120, 240, 120)
            style_run(cell.paragraphs[0].add_run(label), size=9.5, color=GRAY, bold=True)
    doc.add_paragraph()


def activity_1(doc):
    activity_header(doc, 1, "Detetives dos canais digitais", "Escolha o melhor canal para cada situação profissional.")
    student_identification(doc)
    add_callout(doc, "Desafio", "Você faz parte da equipe de comunicação de uma feira escolar. Várias mensagens precisam ser enviadas, mas usar o canal errado pode atrasar o trabalho ou expor informações.", fill="FFF4E5", accent=ORANGE)
    doc.add_heading("Objetivos", level=2)
    add_bullets(doc, [
        "Diferenciar e-mail, chat, videoconferência, documento compartilhado e quadro de tarefas.",
        "Analisar urgência, formalidade, público, registro e segurança.",
        "Justificar uma decisão de comunicação com argumentos claros.",
    ])
    doc.add_heading("Cronograma", level=2)
    add_timeline(doc, [
        ("0-8 min", "Aquecimento", "Em dupla, liste cinco ferramentas usadas para comunicar ou organizar."),
        ("8-15 min", "Critérios", "Leia os cinco critérios e tire dúvidas com o professor."),
        ("15-38 min", "Investigação", "Resolva os seis casos e justifique cada canal escolhido."),
        ("38-50 min", "Confronto", "Compare respostas com outra dupla e defenda duas escolhas."),
        ("50-57 min", "Revisão", "Altere uma resposta depois de ouvir os argumentos do outro grupo."),
        ("57-60 min", "Saída", "Responda individualmente à pergunta final."),
    ])
    doc.add_heading("Os cinco critérios", level=2)
    add_two_column_table(doc, ("Critério", "Pergunta para decidir"), [
        ("Urgência", "A pessoa precisa agir agora, hoje ou nesta semana?"),
        ("Formalidade", "A mensagem representa uma decisão ou compromisso?"),
        ("Complexidade", "O assunto cabe em poucas linhas ou exige conversa?"),
        ("Registro", "Será necessário consultar a informação depois?"),
        ("Segurança", "O canal é autorizado para esse tipo de informação?"),
    ], widths=(1.55, 4.95))
    doc.add_heading("Casos para investigar", level=2)
    cases = [
        "A coordenação precisa aprovar oficialmente o texto do convite.",
        "Um colega não encontrou o link da pasta e precisa dele agora.",
        "Quatro estudantes precisam escrever juntos o regulamento do evento.",
        "A equipe precisa acompanhar quem fará cada tarefa e os prazos.",
        "Existe um desacordo complexo sobre a programação da feira.",
        "Um arquivo com dados pessoais precisa ser entregue ao responsável autorizado.",
    ]
    for idx, case in enumerate(cases, 1):
        p = doc.add_paragraph()
        style_run(p.add_run(f"Caso {idx}. "), bold=True, color=BLUE)
        style_run(p.add_run(case))
        p = doc.add_paragraph()
        style_run(p.add_run("Canal escolhido: "), size=10, color=GRAY, bold=True)
        style_run(p.add_run("________________________________________"), size=10, color="999999")
        p = doc.add_paragraph()
        style_run(p.add_run("Justificativa: "), size=10, color=GRAY, bold=True)
        style_run(p.add_run("____________________________________________________________"), size=10, color="999999")
    doc.add_heading("Pergunta de saída", level=2)
    add_body(doc, "Qual é o maior risco de escolher uma ferramenta apenas porque ela é rápida?")
    response_lines(doc, 3)
    doc.add_heading("Avaliação", level=2)
    add_rubric(doc, [
        ("Escolha do canal", "Canal adequado à situação e aos critérios.", "4"),
        ("Justificativa", "Explica pelo menos dois critérios por caso.", "4"),
        ("Escuta e revisão", "Compara argumentos e revisa uma decisão.", "2"),
    ])
    doc.add_page_break()


def activity_2(doc):
    activity_header(doc, 2, "Missão e-mail profissional", "Transforme uma mensagem confusa em uma comunicação clara e segura.")
    student_identification(doc)
    add_callout(doc, "Situação-problema", "Sua equipe precisa enviar o roteiro de uma apresentação ao professor. Um colega escreveu: 'oi prof ta ai o negocio ve se ta bom preciso resposta urgente'. A missão é reconstruir o e-mail.", fill="FFF4E5", accent=ORANGE)
    doc.add_heading("Objetivos", level=2)
    add_bullets(doc, [
        "Aplicar assunto, saudação, contexto, pedido, prazo, encerramento e assinatura.",
        "Usar tom respeitoso, objetivo e adequado ao mundo do trabalho.",
        "Conferir destinatários, anexos e informações antes do envio.",
    ])
    doc.add_heading("Cronograma", level=2)
    add_timeline(doc, [
        ("0-7 min", "Diagnóstico", "Sublinhe cinco problemas na mensagem original."),
        ("7-15 min", "Modelo", "Revise a estrutura de um e-mail profissional."),
        ("15-35 min", "Produção", "Escreva individualmente a primeira versão."),
        ("35-47 min", "Revisão em dupla", "Use o checklist para revisar o texto de um colega."),
        ("47-56 min", "Versão final", "Aplique as sugestões e prepare a versão definitiva."),
        ("56-60 min", "Fechamento", "Registre a melhoria mais importante realizada."),
    ])
    doc.add_heading("1. Diagnóstico da mensagem", level=2)
    add_body(doc, "Liste cinco problemas encontrados e explique por que cada um prejudica a comunicação.")
    response_lines(doc, 6)
    doc.add_heading("2. Planejamento", level=2)
    add_two_column_table(doc, ("Elemento", "Sua decisão"), [
        ("Assunto específico", ""),
        ("Destinatário", ""),
        ("Contexto", ""),
        ("Pedido ou ação esperada", ""),
        ("Prazo", ""),
        ("Anexo", ""),
        ("Assinatura", ""),
    ], widths=(2.1, 4.4))
    doc.add_heading("3. Primeira versão do e-mail", level=2)
    add_body(doc, "Assunto: ______________________________________________________________________")
    response_lines(doc, 10)
    doc.add_heading("4. Checklist de revisão em dupla", level=2)
    add_bullets(doc, [
        "O assunto permite identificar rapidamente o objetivo?",
        "O texto explica o contexto sem informações desnecessárias?",
        "O pedido e o prazo estão claros?",
        "O tom é respeitoso e não usa cobrança agressiva?",
        "A mensagem menciona e confere o anexo?",
        "Não existem dados pessoais ou destinatários indevidos?",
        "Pontuação, ortografia e assinatura foram revisadas?",
    ])
    doc.add_heading("5. Versão final", level=2)
    response_lines(doc, 11)
    add_body(doc, "Minha melhoria mais importante foi:")
    response_lines(doc, 2)
    doc.add_heading("Avaliação", level=2)
    add_rubric(doc, [
        ("Estrutura", "Contém todos os elementos essenciais.", "3"),
        ("Clareza", "Contexto, pedido e prazo são objetivos.", "3"),
        ("Tom profissional", "Linguagem respeitosa e adequada.", "2"),
        ("Revisão e segurança", "Confere anexo, destinatário e dados.", "2"),
    ])
    doc.add_page_break()


def activity_3(doc):
    activity_header(doc, 3, "Equipe conectada", "Organize arquivos, responsabilidades e tarefas de uma equipe digital.")
    student_identification(doc)
    add_callout(doc, "Desafio", "Uma equipe de quatro estudantes precisa produzir uma campanha sobre uso responsável da internet. Os arquivos estão espalhados, existem cópias diferentes e ninguém sabe quem faz cada parte.", fill="FFF4E5", accent=ORANGE)
    doc.add_heading("Objetivos", level=2)
    add_bullets(doc, [
        "Planejar pastas, nomes de arquivos e permissões de compartilhamento.",
        "Transformar um objetivo em tarefas com responsável, prazo e critério de conclusão.",
        "Comunicar um bloqueio e solicitar ajuda de forma objetiva.",
    ])
    doc.add_heading("Cronograma", level=2)
    add_timeline(doc, [
        ("0-8 min", "Formação da equipe", "Organize grupos de quatro e distribua papéis."),
        ("8-20 min", "Arquitetura de arquivos", "Desenhe pastas, nomes e permissões."),
        ("20-42 min", "Quadro de tarefas", "Cadastre ou planeje oito tarefas completas."),
        ("42-50 min", "Imprevisto", "Resolva um cartão de bloqueio sorteado pelo professor."),
        ("50-57 min", "Apresentação", "Explique o fluxo da equipe em dois minutos."),
        ("57-60 min", "Autoavaliação", "Registre sua contribuição e um aprendizado."),
    ])
    doc.add_heading("Papéis da equipe", level=2)
    add_two_column_table(doc, ("Papel", "Responsabilidade"), [
        ("Organizador", "Define pastas e padrão de nomes."),
        ("Coordenador", "Distribui tarefas e acompanha prazos."),
        ("Revisor", "Confere clareza, segurança e qualidade."),
        ("Relator", "Registra decisões e apresenta o fluxo."),
    ], widths=(1.65, 4.85))
    doc.add_heading("1. Estrutura da pasta compartilhada", level=2)
    add_body(doc, "Desenhe abaixo a pasta principal e pelo menos quatro subpastas. Indique quem pode visualizar, comentar ou editar.")
    response_lines(doc, 8)
    doc.add_heading("2. Padrão de nomes", level=2)
    add_body(doc, "Crie nomes claros para os arquivos: roteiro, pesquisa, imagens, apresentação e entrega final.")
    response_lines(doc, 6)
    doc.add_heading("3. Quadro de tarefas", level=2)
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = (1.7, 1.05, 0.85, 1.1, 1.8)
    for idx, text in enumerate(("Tarefa", "Responsável", "Prazo", "Status", "Concluída quando...")):
        cell = table.rows[0].cells[idx]
        cell.width = Inches(widths[idx])
        set_cell_shading(cell, BLUE)
        set_cell_margins(cell, 80, 80, 80, 80)
        style_run(cell.paragraphs[0].add_run(text), size=8.5, color="FFFFFF", bold=True)
    for _ in range(8):
        cells = table.add_row().cells
        for idx, cell in enumerate(cells):
            cell.width = Inches(widths[idx])
            set_cell_margins(cell, 80, 80, 280, 80)
    doc.add_paragraph()
    doc.add_heading("4. Cartão de bloqueio", level=2)
    add_body(doc, "Problema recebido: ______________________________________________________________")
    add_body(doc, "Mensagem de ajuda que a equipe enviaria:")
    response_lines(doc, 4)
    doc.add_heading("Autoavaliação individual", level=2)
    add_body(doc, "Minha principal contribuição foi:")
    response_lines(doc, 2)
    add_body(doc, "Aprendi que uma equipe digital precisa:")
    response_lines(doc, 2)
    doc.add_heading("Avaliação", level=2)
    add_rubric(doc, [
        ("Organização", "Pastas, nomes e permissões coerentes.", "3"),
        ("Qualidade das tarefas", "Ação, responsável, prazo e conclusão claros.", "3"),
        ("Comunicação do bloqueio", "Contexto e pedido de ajuda objetivos.", "2"),
        ("Colaboração", "Papéis e participação equilibrados.", "2"),
    ])
    doc.add_page_break()


def activity_4(doc):
    activity_header(doc, 4, "Central de segurança digital", "Analise mensagens, reconheça riscos e responda com netiqueta.")
    student_identification(doc)
    add_callout(doc, "Missão", "Você faz parte de uma central que recebeu quatro mensagens. A equipe deve identificar phishing, exposição de dados, falhas de netiqueta e compartilhamento inseguro.", fill="FDECEA", accent=RED)
    doc.add_heading("Objetivos", level=2)
    add_bullets(doc, [
        "Reconhecer sinais de phishing e pedidos suspeitos.",
        "Aplicar netiqueta ao responder conflitos e solicitações.",
        "Proteger senhas, dados pessoais, links e permissões.",
        "Saber quando interromper a ação e pedir ajuda a um responsável.",
    ])
    doc.add_heading("Cronograma", level=2)
    add_timeline(doc, [
        ("0-8 min", "Radar de riscos", "Liste sinais que tornam uma mensagem suspeita."),
        ("8-15 min", "Protocolo", "Leia os passos Parar, Verificar, Proteger e Comunicar."),
        ("15-40 min", "Análise", "Em trio, investigue os quatro casos simulados."),
        ("40-50 min", "Resposta segura", "Reescreva uma mensagem e crie orientação para a equipe."),
        ("50-57 min", "Coletiva", "Compare classificações e ajuste uma decisão."),
        ("57-60 min", "Compromisso", "Registre uma atitude que adotará daqui em diante."),
    ])
    doc.add_heading("Protocolo PVPC", level=2)
    add_steps(doc, [
        "Parar: não clicar, responder ou encaminhar por impulso.",
        "Verificar: conferir remetente, domínio, link, pedido e contexto.",
        "Proteger: não informar senha, código ou dado pessoal; revisar permissões.",
        "Comunicar: avisar professor, responsável ou canal oficial quando houver risco.",
    ])
    doc.add_heading("Casos simulados", level=2)
    cases = [
        ("Mensagem A", "URGENTE! Sua conta será apagada. Clique em encurta.link/conta e confirme sua senha em 5 minutos."),
        ("Mensagem B", "Um colega publica no grupo: 'Você sempre atrasa tudo. Resolve isso logo!!!'"),
        ("Mensagem C", "Uma planilha com nome completo, telefone e endereço dos participantes está configurada como 'qualquer pessoa com o link'."),
        ("Mensagem D", "Um arquivo inesperado chamado FotosEvento.exe chega de um endereço parecido, mas diferente do oficial."),
    ]
    for label, text in cases:
        add_callout(doc, label, text, fill="F3F5F7", accent=BLUE)
        add_body(doc, "Classificação: [  ] seguro   [  ] inadequado   [  ] suspeito   [  ] fraude provável")
        add_body(doc, "Sinais encontrados e ação recomendada:")
        response_lines(doc, 3)
    doc.add_heading("Resposta segura e respeitosa", level=2)
    add_body(doc, "Reescreva a Mensagem B usando fato, necessidade, pedido claro e prazo.")
    response_lines(doc, 5)
    doc.add_heading("Aviso para a equipe", level=2)
    add_body(doc, "Escreva um aviso curto com três orientações para evitar os riscos analisados.")
    response_lines(doc, 5)
    doc.add_heading("Compromisso individual", level=2)
    add_body(doc, "A partir de hoje, antes de clicar, compartilhar ou responder, eu vou:")
    response_lines(doc, 3)
    doc.add_heading("Avaliação", level=2)
    add_rubric(doc, [
        ("Identificação de riscos", "Reconhece sinais de phishing, exposição e conflito.", "4"),
        ("Ação segura", "Aplica o protocolo PVPC e busca ajuda quando necessário.", "3"),
        ("Netiqueta", "Reescreve a mensagem com respeito e objetividade.", "2"),
        ("Participação", "Contribui com a investigação do trio.", "1"),
    ])


def build():
    doc = Document()
    configure_document(doc)
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.clear()
    style_run(header.add_run("SENAI | Atividades - Ferramentas Digitais para Comunicação"), size=8.5, color=GRAY, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(85)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_run(p.add_run("CADERNO DE ATIVIDADES"), size=12, color=ORANGE, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(8)
    style_run(p.add_run("Ferramentas Digitais\npara Comunicação"), size=28, color=BLUE, bold=True, font="Aptos Display")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_run(p.add_run("4 atividades práticas de 1 hora | Estudantes de 14 anos"), size=12.5, color=GRAY, italic=True)
    add_callout(doc, "Competências trabalhadas", "Escolha de canais, e-mail profissional, colaboração em nuvem, gestão de tarefas, netiqueta, segurança e privacidade.", fill=LIGHT_BLUE)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(75)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_run(p.add_run("Programa Rio do Sul Mais Tech\nSENAI / Prefeitura Municipal de Rio do Sul"), size=11, color=BLUE, bold=True)
    doc.add_page_break()

    doc.add_heading("Orientações gerais", level=1)
    add_body(doc, "As atividades foram planejadas para aulas de 60 minutos e podem ser realizadas em papel, em ambiente digital autorizado ou de forma híbrida. Não utilize contas pessoais, senhas verdadeiras, dados reais ou links suspeitos.")
    add_bullets(doc, [
        "Organização sugerida: duplas, trios e equipes de quatro estudantes.",
        "Recursos: apostila, projetor, folhas, canetas e ferramentas autorizadas pela escola.",
        "Avaliação: 10 pontos por atividade, com rubrica ao final de cada encontro.",
        "Portfólio: guardar as quatro entregas para acompanhar a evolução do estudante.",
        "Segurança: todos os exemplos de phishing devem ser simulados e não clicáveis.",
    ])
    add_two_column_table(doc, ("Atividade", "Produto final"), [
        ("1. Detetives dos canais", "Mapa de canais com justificativas."),
        ("2. Missão e-mail", "E-mail profissional revisado."),
        ("3. Equipe conectada", "Estrutura de arquivos e quadro de tarefas."),
        ("4. Central de segurança", "Análise de riscos e resposta segura."),
    ], widths=(2.25, 4.25))
    doc.add_page_break()

    activity_1(doc)
    activity_2(doc)
    activity_3(doc)
    activity_4(doc)

    props = doc.core_properties
    props.title = "Atividades Ferramentas Digitais para Comunicação"
    props.subject = "Quatro atividades de 60 minutos para estudantes de 14 anos"
    props.author = "SENAI - Programa Rio do Sul Mais Tech"
    props.keywords = "atividades, comunicação digital, e-mail, colaboração, segurança"

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
