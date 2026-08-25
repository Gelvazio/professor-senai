from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
TARGET = ROOT / "sistema" / "FICHA-PRODUTO-MAIS-TECH" / "INTRODUCAO_COMUNICACAO_ORAL_ESCRITA" / "04 - FERRAMENTAS-DIGITAIS" / "AtividadesFerramentasDigitaisFacil.docx"

REPLACEMENTS = (
    ("ATIVIDADE FÁCIL", "ATIVIDADE"),
    ("Atividade Fácil", "Atividade"),
    ("Atividade fácil", "Atividade"),
    ("atividade fácil", "atividade"),
    ("FÁCIL", ""),
    ("Fácil", "Consegui"),
    ("fácil", ""),
    ("FACIL", ""),
    ("Facil", "Consegui"),
    ("facil", ""),
)


def replace_text(text):
    for old, new in REPLACEMENTS:
        text = text.replace(old, new)
    return " ".join(text.split()) if text.strip() else text


def replace_paragraph(paragraph):
    for run in paragraph.runs:
        run.text = replace_text(run.text)


def all_paragraphs(doc):
    for paragraph in doc.paragraphs:
        yield paragraph
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs
                for nested_table in cell.tables:
                    for nested_row in nested_table.rows:
                        for nested_cell in nested_row.cells:
                            yield from nested_cell.paragraphs
    for section in doc.sections:
        yield from section.header.paragraphs
        yield from section.footer.paragraphs


def edit_document():
    if not TARGET.exists():
        raise FileNotFoundError(TARGET)

    doc = Document(TARGET)
    for paragraph in all_paragraphs(doc):
        replace_paragraph(paragraph)

    props = doc.core_properties
    props.title = replace_text(props.title or "")
    props.subject = replace_text(props.subject or "")
    props.keywords = replace_text(props.keywords or "")
    props.comments = replace_text(props.comments or "")

    doc.save(TARGET)
    print(TARGET)


if __name__ == "__main__":
    edit_document()
