from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path(
    r"C:\fontes\professor-senai\sistema\FICHA-PRODUTO-MAIS-TECH"
    r"\REFORCO_LINGUAGENS\TEXTO-REFORCO-25-08-2026.docx"
)

TITLE = "A IMPORTÂNCIA DA LINGUAGEM EM NOSSA VIDA"

PARAGRAPHS = [
    "A linguagem está presente em quase todos os momentos de nossa vida. Nós a usamos para conversar com a família, estudar, fazer perguntas, contar histórias, compreender avisos e expressar sentimentos. Também usamos a linguagem quando escrevemos uma mensagem, lemos uma notícia ou apresentamos um trabalho. Por isso, aprender a ler, escrever, ouvir e falar com atenção é importante não apenas para ter boas notas, mas também para participar da sociedade com segurança e responsabilidade.",
    "Ler é muito mais do que reconhecer palavras. Uma boa leitura exige que a pessoa compreenda o sentido geral do texto e perceba seus detalhes. Antes de começar, podemos observar o título e pensar sobre o assunto que será apresentado. Durante a leitura, devemos identificar o tema, a ideia principal e as informações que ajudam a explicar essa ideia. Quando encontramos uma palavra desconhecida, podemos tentar descobrir seu significado pelo contexto ou consultar um dicionário.",
    "Todo texto é produzido com alguma intenção. Uma notícia procura informar; uma propaganda tenta convencer; uma receita ensina como preparar algo; e uma história pode divertir e provocar reflexões. Reconhecer essa intenção nos ajuda a interpretar melhor o que lemos. Também é importante perguntar quem escreveu o texto, para quem ele foi escrito e em qual situação circula. Essas perguntas desenvolvem o pensamento crítico e evitam que aceitemos qualquer informação sem refletir.",
    "Na internet, o cuidado precisa ser ainda maior. Nem tudo o que aparece em vídeos, mensagens e redes sociais é verdadeiro. Antes de compartilhar uma informação, devemos verificar a fonte, comparar o conteúdo com outros materiais e observar a data de publicação. Uma mensagem pode parecer convincente e, mesmo assim, apresentar erros ou tentar enganar o leitor. Ser um leitor atento significa investigar, fazer perguntas e formar uma opinião com base em informações confiáveis.",
    "A escrita também precisa ser organizada. Um texto claro apresenta ideias que se relacionam e seguem uma ordem compreensível. O primeiro parágrafo pode introduzir o assunto; os parágrafos seguintes desenvolvem as informações; e o último apresenta uma conclusão. Cada parágrafo deve tratar de uma ideia principal. Para unir as partes, podemos usar palavras como também, porém, porque, por isso, depois e finalmente. Essas palavras colaboram com a coesão e tornam a leitura mais fácil.",
    "A pontuação ajuda o leitor a entender o texto. O ponto final encerra uma ideia completa. A vírgula pode separar elementos de uma lista ou marcar pequenas pausas. O ponto de interrogação aparece em perguntas, enquanto o ponto de exclamação indica surpresa, alegria, ordem ou outra emoção intensa. Usar os sinais corretamente evita confusões. A frase “Vamos comer, crianças” não tem o mesmo sentido de “Vamos comer crianças”. Uma simples vírgula pode mudar toda a mensagem.",
    "Outro cuidado necessário é a concordância. As palavras de uma frase precisam combinar entre si. Dizemos “os estudantes atentos participaram” porque o artigo, o substantivo e o adjetivo estão no plural. Também escrevemos “a turma apresentou o trabalho” porque o verbo concorda com o núcleo do sujeito, que está no singular. Revisar essas combinações melhora a escrita e ajuda o leitor a compreender exatamente o que queremos comunicar.",
    "Comunicar-se bem também envolve saber ouvir. Em uma conversa ou debate, cada pessoa deve esperar sua vez, prestar atenção e respeitar opiniões diferentes. Discordar não significa atacar ou humilhar alguém. Podemos apresentar argumentos, explicar nossos motivos e fazer perguntas de maneira educada. Quando trabalhamos em grupo, a comunicação clara permite dividir tarefas, cumprir prazos e resolver problemas com mais facilidade. Assim, todos podem contribuir para o resultado.",
    "Para desenvolver autonomia nos estudos, é útil criar uma rotina. Podemos escolher um local tranquilo, separar os materiais e definir um tempo para cada atividade. Fazer anotações, destacar ideias importantes e produzir pequenos resumos ajuda a guardar o conteúdo. Quando uma dúvida aparece, devemos reler, pesquisar ou pedir orientação. Estudar com autonomia não significa fazer tudo sozinho, mas assumir responsabilidade pela própria aprendizagem e procurar ajuda quando necessário.",
    "O domínio da linguagem cresce com a prática. Quanto mais lemos, ampliamos nosso vocabulário e conhecemos diferentes formas de organizar as ideias. Quanto mais escrevemos e revisamos, percebemos nossos avanços e corrigimos dificuldades. Ao falar e ouvir com respeito, construímos relações melhores. Portanto, aprender Linguagens é aprender a compreender o mundo, comunicar pensamentos e participar de decisões. Cada leitura, cada texto e cada conversa é uma oportunidade de aprender e de fazer nossa voz ser ouvida.",
]


def set_cell_free_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, end])


def build_document():
    doc = Document()
    section = doc.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(12)
    title.paragraph_format.keep_with_next = True
    run = title.add_run(TITLE)
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(46, 116, 181)

    for text in PARAGRAPHS:
        paragraph = doc.add_paragraph(text)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.first_line_indent = Inches(0.3)
        paragraph.paragraph_format.widow_control = True

    footer = section.footer.paragraphs[0]
    footer.paragraph_format.space_before = Pt(0)
    footer.paragraph_format.space_after = Pt(0)
    set_cell_free_page_number(footer)

    doc.core_properties.title = TITLE.title()
    doc.core_properties.subject = "Reforço de Linguagens para alunos de 13 anos"
    doc.core_properties.author = "SENAI — Rio do Sul Mais Tech"
    doc.core_properties.keywords = "linguagens, leitura, escrita, comunicação, reforço"

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)


if __name__ == "__main__":
    build_document()
