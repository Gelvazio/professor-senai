#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
Gerador DOCX inteligente a partir de Markdown
"""

from pathlib import Path
import re

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    print("❌ python-docx não instalado")


def determine_lines_needed(text):
    """Determina a quantidade de linhas necessárias"""

    text_lower = text.lower()

    # Análise e reflexão
    if "mínimo" in text_lower and "linha" in text_lower:
        match = re.search(r'mínimo\s+(\d+)\s+linha', text_lower)
        if match:
            num = int(match.group(1))
            return num

    # Perguntas discursivas
    discursive_keywords = ["como você pensa", "como", "por quê", "porque",
                          "explique", "analise", "descreva", "qual foi a inovação"]

    if any(x in text_lower for x in discursive_keywords):
        if text.strip().endswith("?"):
            if "como você pensa" in text_lower or "inovação" in text_lower:
                return 8
            elif len(text) > 100:
                return 6
            else:
                return 4

    # Respostas simples
    if "resposta:" in text_lower and "(" not in text:
        return 3

    # Campos com underscores
    if "_" * 10 in text:
        return 3

    return 0


def markdown_to_docx(md_file, docx_file):
    """Converte Markdown para DOCX"""

    if not HAS_DOCX:
        print("❌ python-docx não disponível")
        return False

    try:
        with open(md_file, 'r', encoding='utf-8') as f:
            content = f.read()

        doc = Document()

        # Cores padrão
        color_primary = RGBColor(27, 67, 50)  # #1B4332
        color_secondary = RGBColor(45, 106, 79)  # #2D6A4F

        lines = content.split('\n')
        i = 0

        while i < len(lines):
            line = lines[i].rstrip()
            i += 1

            if not line.strip():
                continue

            # Títulos principais
            if line.startswith('# '):
                text = line[2:].strip()
                p = doc.add_paragraph(text)
                p.style = 'Heading 1'
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                for run in p.runs:
                    run.font.color.rgb = color_primary
                    run.font.size = Pt(16)
                    run.font.bold = True

            # Heading 2
            elif line.startswith('## '):
                text = line[3:].strip()
                p = doc.add_paragraph(text)
                p.style = 'Heading 2'
                for run in p.runs:
                    run.font.color.rgb = color_secondary
                    run.font.size = Pt(12)
                    run.font.bold = True

            # Heading 3
            elif line.startswith('### '):
                text = line[4:].strip()
                p = doc.add_paragraph(text)
                p.style = 'Heading 3'
                for run in p.runs:
                    run.font.color.rgb = color_secondary
                    run.font.size = Pt(11)
                    run.font.bold = True

            # Tabelas
            elif line.startswith('|'):
                table_rows = []
                while i < len(lines) and lines[i].startswith('|'):
                    row = lines[i].strip()
                    if '---' not in row:
                        cells = [cell.strip() for cell in row.split('|')[1:-1]]
                        table_rows.append(cells)
                    i += 1

                if table_rows:
                    cols = len(table_rows[0])
                    rows = len(table_rows)
                    table = doc.add_table(rows=rows, cols=cols)
                    table.style = 'Light Grid Accent 1'

                    for i_row, row_data in enumerate(table_rows):
                        for i_cell, cell_text in enumerate(row_data):
                            cell = table.rows[i_row].cells[i_cell]
                            cell.text = cell_text

                            # Cabeçalho
                            if i_row == 0:
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        run.font.bold = True
                                        run.font.color.rgb = color_secondary
                continue

            # Listas
            elif line.startswith('- '):
                text = line[2:].strip()
                p = doc.add_paragraph(text, style='List Bullet')

            elif re.match(r'^\d+\.\s', line):
                match = re.match(r'^(\d+)\.\s(.+)$', line)
                if match:
                    text = match.group(2)
                    p = doc.add_paragraph(text, style='List Number')

            # Parágrafos com perguntas
            else:
                if line.strip():
                    # Remover formatação perigosa
                    text = line.replace('**', '').replace('__', '').replace('_', '')
                    text = text.replace('*', '')

                    p = doc.add_paragraph(text)

                    # Se é pergunta, adicionar espaço para resposta
                    if text.strip().endswith('?'):
                        num_lines = determine_lines_needed(text)
                        if num_lines > 0:
                            # Adicionar linhas em branco
                            for _ in range(num_lines):
                                doc.add_paragraph('_' * 100)

        # Salvar DOCX
        doc.save(docx_file)
        print(f"✅ DOCX criado: {docx_file}")
        return True

    except Exception as e:
        print(f"❌ Erro ao criar DOCX: {e}")
        import traceback
        traceback.print_exc()
        return False


def main():
    """Função principal"""

    script_dir = Path(__file__).parent
    md_files = ['ATIVIDADE-EPISODIO-01.md', 'ATIVIDADE-EPISODIO-02.md']

    print("=" * 60)
    print("GERADOR DOCX - Saga dos Computadores")
    print("=" * 60)

    for md_file in md_files:
        md_path = script_dir / md_file
        docx_path = script_dir / md_file.replace('.md', '.docx')

        if md_path.exists():
            print(f"\n📄 Gerando: {md_file}...")
            markdown_to_docx(str(md_path), str(docx_path))
        else:
            print(f"⚠️  Arquivo não encontrado: {md_file}")

    print("\n" + "=" * 60)
    print("✅ CONCLUÍDO!")
    print("=" * 60)


if __name__ == '__main__':
    main()
