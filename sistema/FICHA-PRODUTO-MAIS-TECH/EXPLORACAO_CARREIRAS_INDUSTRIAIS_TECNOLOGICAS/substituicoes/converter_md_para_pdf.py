#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
Script para converter Markdown para PDF
Cria PDFs a partir dos arquivos ATIVIDADE-EPISODIO-*.md
"""

import os
from pathlib import Path
import re

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    print("⚠️ python-docx não instalado. Instale com: pip install python-docx")

try:
    from docx2pdf.converter import convert
    HAS_DOCX2PDF = True
except ImportError:
    HAS_DOCX2PDF = False
    print("⚠️ docx2pdf não instalado. Instale com: pip install docx2pdf")


def markdown_to_docx(md_file, docx_file):
    """Converte Markdown para DOCX"""

    if not HAS_DOCX:
        print("❌ python-docx não disponível")
        return False

    try:
        # Ler arquivo Markdown
        with open(md_file, 'r', encoding='utf-8') as f:
            content = f.read()

        # Criar documento Word
        doc = Document()

        # Processar conteúdo linha por linha
        lines = content.split('\n')
        in_table = False
        table_data = []

        for line in lines:
            line = line.rstrip()

            # Pular linhas vazias múltiplas
            if not line:
                if not in_table:
                    doc.add_paragraph()
                continue

            # Títulos
            if line.startswith('# '):
                p = doc.add_paragraph(line[2:], style='Heading 1')
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            elif line.startswith('## '):
                doc.add_paragraph(line[3:], style='Heading 2')

            elif line.startswith('### '):
                doc.add_paragraph(line[4:], style='Heading 3')

            # Tabelas (simples)
            elif line.startswith('|'):
                table_data.append(line)

                # Detectar fim da tabela
                if table_data and not line.startswith('|'):
                    # Processar tabela
                    if len(table_data) > 2:
                        rows = [row.split('|')[1:-1] for row in table_data if row.startswith('|')]
                        rows = [[cell.strip() for cell in row] for row in rows]

                        if rows:
                            table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                            table.style = 'Light Grid Accent 1'

                            for i, row in enumerate(rows):
                                for j, cell in enumerate(row):
                                    table.rows[i].cells[j].text = cell

                    table_data = []
                    in_table = False
                else:
                    in_table = True

            # Listas
            elif line.startswith('- '):
                p = doc.add_paragraph(line[2:], style='List Bullet')

            elif line.startswith('* '):
                p = doc.add_paragraph(line[2:], style='List Bullet')

            # Números de lista
            elif re.match(r'^\d+\.\s', line):
                match = re.match(r'^(\d+)\.\s(.+)$', line)
                if match:
                    p = doc.add_paragraph(match.group(2), style='List Number')

            # Código
            elif line.startswith('```'):
                continue

            # Negrito e itálico
            else:
                p = doc.add_paragraph()

                # Processar formatação
                text = line
                text = re.sub(r'\*\*(.+?)\*\*', r'__\1__', text)  # **bold** → __bold__
                text = re.sub(r'\*(.+?)\*', r'_\1_', text)  # *italic* → _italic_

                # Adicionar com formatação
                if '**' in line or '__' in text:
                    parts = re.split(r'(__[^_]+__)', text)
                    for part in parts:
                        if part.startswith('__') and part.endswith('__'):
                            run = p.add_run(part[2:-2])
                            run.bold = True
                        else:
                            p.add_run(part)
                else:
                    p.add_run(text)

        # Processar última tabela se houver
        if table_data and len(table_data) > 2:
            rows = [row.split('|')[1:-1] for row in table_data if row.startswith('|')]
            rows = [[cell.strip() for cell in row] for row in rows]

            if rows:
                table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                table.style = 'Light Grid Accent 1'

                for i, row in enumerate(rows):
                    for j, cell in enumerate(row):
                        table.rows[i].cells[j].text = cell

        # Salvar DOCX
        doc.save(docx_file)
        print(f"✅ DOCX criado: {docx_file}")
        return True

    except Exception as e:
        print(f"❌ Erro ao converter {md_file}: {e}")
        return False


def docx_to_pdf(docx_file, pdf_file):
    """Converte DOCX para PDF"""

    if not HAS_DOCX2PDF:
        print("❌ docx2pdf não disponível. Instale com: pip install docx2pdf")
        return False

    try:
        convert(docx_file, pdf_file)
        print(f"✅ PDF criado: {pdf_file}")
        return True
    except Exception as e:
        print(f"❌ Erro ao converter {docx_file} para PDF: {e}")
        return False


def main():
    """Função principal"""

    # Diretório de trabalho
    script_dir = Path(__file__).parent

    # Arquivos Markdown para converter
    md_files = [
        'ATIVIDADE-EPISODIO-01.md',
        'ATIVIDADE-EPISODIO-02.md'
    ]

    print("=" * 60)
    print("CONVERSOR MARKDOWN → DOCX → PDF")
    print("=" * 60)

    for md_filename in md_files:
        md_path = script_dir / md_filename
        docx_path = script_dir / md_filename.replace('.md', '.docx')
        pdf_path = script_dir / md_filename.replace('.md', '.pdf')

        if not md_path.exists():
            print(f"⚠️  {md_filename} não encontrado")
            continue

        print(f"\n📄 Processando {md_filename}...")

        # Converter Markdown → DOCX
        if markdown_to_docx(str(md_path), str(docx_path)):
            # Converter DOCX → PDF
            if HAS_DOCX2PDF:
                docx_to_pdf(str(docx_path), str(pdf_path))
            else:
                print(f"⚠️  Para converter para PDF, instale: pip install docx2pdf")
        else:
            print(f"❌ Falha ao processar {md_filename}")

    print("\n" + "=" * 60)
    print("CONVERSÃO CONCLUÍDA!")
    print("=" * 60)


if __name__ == '__main__':
    main()
