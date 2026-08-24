from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "INTRODUCAO_COMUNICACAO_ORAL_ESCRITA" / "AVALIACOES" / "RECUPERACAO-COMUNICACAO_ORAL_ESCRITA.docx"
PETROLEO = "004D5C"
LARANJA = "FF6B35"
AZUL_CLARO = "EAF4F6"
CINZA = "5F6B73"


def style_run(run, size=10.5, bold=False, color="000000", italic=False):
    run.font.name = "Arial"
    fonts = run._element.get_or_add_rPr().rFonts
    fonts.set(qn("w:ascii"), "Arial")
    fonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def set_fill(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), color)


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_margins(cell, top=90, start=120, bottom=90, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            set_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(5)
    style_run(p.add_run(text), 15 if level == 1 else 12, True, PETROLEO)


def body(doc, text, bold_start=None, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    if bold_start and text.startswith(bold_start):
        prefix, rest = text.split(":", 1)
        style_run(p.add_run(prefix + ":"), bold=True, color=PETROLEO)
        style_run(p.add_run(rest), italic=italic)
    else:
        style_run(p.add_run(text), italic=italic)


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(3)
    style_run(p.add_run(text))


def numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(4)
    style_run(p.add_run(text))


def callout(doc, label, text):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    table_geometry(table, [9360])
    set_fill(table.cell(0, 0), "FFF2EB")
    p = table.cell(0, 0).paragraphs[0]
    style_run(p.add_run(label + " "), 10, True, LARANJA)
    style_run(p.add_run(text), 10)


def add_checklist(doc, items):
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        style_run(p.add_run("☐ "), 11, True, LARANJA)
        style_run(p.add_run(item))


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.7)
section.bottom_margin = Inches(0.7)
section.left_margin = Inches(0.75)
section.right_margin = Inches(0.75)
section.header_distance = Inches(0.3)
section.footer_distance = Inches(0.3)

normal = doc.styles["Normal"]
normal.font.name = "Arial"
normal.font.size = Pt(10.5)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.15
for name, size in (("Heading 1", 15), ("Heading 2", 12), ("Heading 3", 11)):
    style = doc.styles[name]
    style.font.name = "Arial"
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor.from_string(PETROLEO)

header = section.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.CENTER
style_run(header.add_run("RIO DO SUL MAIS TECH · SENAI · Prefeitura Municipal de Rio do Sul"), 8.5, True, PETROLEO)
footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
style_run(footer.add_run("Recuperação · Comunicação Oral e Escrita para o Mundo do Trabalho"), 8, color=CINZA)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(8)
p.paragraph_format.space_after = Pt(2)
style_run(p.add_run("RECUPERAÇÃO DA AVALIAÇÃO PRÁTICA"), 21, True, PETROLEO)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(12)
style_run(p.add_run("Portfólio Individual de Comunicação Profissional"), 13, True, LARANJA)

info = doc.add_table(rows=2, cols=2)
info.style = "Table Grid"
table_geometry(info, [4680, 4680])
for row, labels in zip(info.rows, (("Aluno(a):", "Turma:"), ("Data de entrega:", "Professor:"))):
    for cell, label in zip(row.cells, labels):
        set_fill(cell, AZUL_CLARO)
        style_run(cell.paragraphs[0].add_run(f"{label} ______________________________"), 10, True, PETROLEO)

heading(doc, "1. Objetivo da recuperação")
body(doc, "Demonstrar, individualmente, domínio das capacidades desenvolvidas em toda a Unidade Curricular por meio de um material completo que reúna análise da comunicação, produção de documentos profissionais, uso responsável de ferramentas digitais e apresentação oral.")
callout(doc, "Entrega obrigatória:", "um portfólio individual organizado, digitado e revisado, acompanhado de apresentação oral de 5 a 7 minutos. Valor total: 10,0 pontos. Nota mínima: 7,0.")

heading(doc, "2. Situação integradora")
body(doc, "Você foi escolhido para atuar como assistente de comunicação da empresa fictícia Conecta Jovem, que oferece oficinas de tecnologia para escolas. A empresa está enfrentando falhas de comunicação: mensagens confusas, atrasos, atendimento inadequado, reuniões sem registro e riscos no uso das ferramentas digitais. Sua missão é criar um Plano Completo de Comunicação Profissional para melhorar esse cenário.")

heading(doc, "3. Formato e organização da entrega")
body(doc, "O material deverá ser produzido em editor de texto (Google Docs, Microsoft Word ou equivalente), salvo em PDF para entrega e organizado nesta ordem:")
for item in (
    "Capa com nome do aluno, turma, título do trabalho e data;",
    "Sumário com as nove partes do portfólio;",
    "Textos autorais, completos e revisados;",
    "Títulos e subtítulos claros;",
    "Fontes consultadas ao final;",
    "Apresentação oral com apoio visual opcional."
):
    bullet(doc, item)
body(doc, "Extensão sugerida: 8 a 12 páginas, além da capa e do sumário. Todos os documentos devem ser elaborados pelo próprio aluno.")

doc.add_page_break()
heading(doc, "4. Partes obrigatórias do portfólio")

heading(doc, "PARTE 1 — Fundamentos e processo comunicativo", 2)
body(doc, "Produza um texto de 15 a 20 linhas explicando a importância da comunicação no mundo do trabalho. Em seguida, represente uma situação profissional e identifique:")
for item in ("emissor;", "receptor;", "mensagem;", "canal;", "código;", "ruído;", "feedback."):
    bullet(doc, item)
body(doc, "Inclua uma conclusão explicando como o feedback poderia evitar ou corrigir o ruído identificado.")

heading(doc, "PARTE 2 — Barreiras, linguagem e adequação ao contexto", 2)
body(doc, "Crie duas versões da mesma mensagem: uma informal, adequada a um amigo, e outra formal, adequada a uma empresa. Depois:")
numbered(doc, "Explique a diferença entre comunicação verbal e não verbal.")
numbered(doc, "Apresente três barreiras de comunicação e uma solução para cada uma.")
numbered(doc, "Descreva como postura, expressão facial, gestos e tom de voz podem alterar o sentido de uma mensagem.")

heading(doc, "PARTE 3 — Escuta ativa, empatia e comunicação assertiva", 2)
body(doc, "Analise a situação: um cliente está irritado porque recebeu uma informação incorreta e fala em tom elevado com o atendente.")
numbered(doc, "Escreva uma resposta passiva, uma agressiva e uma assertiva.")
numbered(doc, "Explique por que a resposta assertiva é a mais adequada.")
numbered(doc, "Crie um roteiro de atendimento com saudação, escuta ativa, confirmação do problema, proposta de solução e encerramento cortês.")

heading(doc, "PARTE 4 — E-mail corporativo", 2)
body(doc, "Redija um e-mail para uma escola informando que a oficina marcada para terça-feira, às 14h, foi transferida para quarta-feira, às 15h, no mesmo local. Solicite confirmação de recebimento.")
add_checklist(doc, ("Assunto objetivo", "Saudação", "Corpo claro e completo", "Orientação de confirmação", "Despedida", "Assinatura profissional"))

heading(doc, "PARTE 5 — Comunicado e memorando", 2)
body(doc, "Produza os dois documentos abaixo:")
numbered(doc, "Comunicado aos participantes sobre novas regras de uso do laboratório: pontualidade, cuidado com equipamentos, alimentação e organização.")
numbered(doc, "Memorando interno à equipe solicitando revisão dos materiais da próxima oficina, com objetivo, responsáveis e prazo.")

heading(doc, "PARTE 6 — Ata de reunião", 2)
body(doc, "Elabore a ata de uma reunião da equipe Conecta Jovem. A reunião ocorreu em 20/08/2026, das 14h às 15h, com quatro participantes e discutiu reclamações dos clientes, revisão das mensagens e divisão de responsabilidades. A ata deve registrar abertura, participantes, pauta, decisões, responsáveis, prazos e encerramento.")

heading(doc, "PARTE 7 — Relato de ocorrência", 2)
body(doc, "Redija um relato factual sobre a seguinte ocorrência: durante uma oficina, um estudante recebeu por mensagem um link falso que solicitava senha. Ele clicou no link, mas avisou o professor antes de preencher os dados. Registre data, horário, local, envolvidos, fatos observados, providências adotadas e encaminhamento, sem opiniões ou julgamentos.")

heading(doc, "PARTE 8 — Ferramentas digitais, netiqueta e segurança", 2)
body(doc, "Crie um guia prático intitulado Boas Práticas de Comunicação Digital da Conecta Jovem. O guia deve apresentar:")
for item in (
    "uso adequado do e-mail corporativo;",
    "regras de netiqueta em chats, reuniões on-line e documentos compartilhados;",
    "uma proposta de uso do Google Workspace ou Microsoft 365 para comunicação e colaboração;",
    "um aplicativo para organizar tarefas, com explicação de como será utilizado;",
    "cinco cuidados com senhas;",
    "cinco sinais de phishing;",
    "cuidados com privacidade e compartilhamento de dados."
):
    bullet(doc, item)

heading(doc, "PARTE 9 — Reunião e apresentação oral", 2)
body(doc, "Prepare uma apresentação individual de 5 a 7 minutos para defender seu Plano Completo de Comunicação Profissional. A apresentação deverá:")
for item in (
    "apresentar o problema e a solução proposta;",
    "explicar pelo menos um documento profissional produzido;",
    "demonstrar clareza, objetividade e sequência lógica;",
    "usar postura, contato visual, tom de voz e gestos adequados;",
    "simular uma breve fala de reunião ou atendimento ao cliente;",
    "responder a uma pergunta do professor usando escuta ativa e assertividade."
):
    bullet(doc, item)

doc.add_page_break()
heading(doc, "5. Checklist antes da entrega")
add_checklist(doc, (
    "Incluí capa e sumário.",
    "Completei as nove partes obrigatórias.",
    "Produzi e-mail, comunicado, memorando, ata e relato de ocorrência.",
    "Diferenciei comunicação formal e informal, verbal e não verbal.",
    "Expliquei ruídos, barreiras, feedback, escuta ativa, empatia e assertividade.",
    "Incluí atendimento ao cliente, reunião e apresentação oral.",
    "Abordei ferramentas colaborativas, gestão de tarefas, netiqueta, phishing, senhas e privacidade.",
    "Revisei clareza, objetividade, coesão, coerência, paragrafação, pontuação, concordância e ortografia.",
    "Listei as fontes consultadas.",
    "Preparei minha apresentação de 5 a 7 minutos."
))

heading(doc, "6. Rubrica de avaliação")
rubric = doc.add_table(rows=1, cols=4)
rubric.style = "Table Grid"
widths = [2800, 4160, 1200, 1200]
table_geometry(rubric, widths)
for cell, label in zip(rubric.rows[0].cells, ("Critério", "Evidências esperadas", "Valor", "Nota")):
    set_fill(cell, PETROLEO)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_run(p.add_run(label), 9.3, True, "FFFFFF")

criteria = (
    ("Cobertura da ementa", "As nove partes foram entregues e contemplam todos os temas solicitados.", "2,0"),
    ("Fundamentos e análise", "Aplica corretamente processo comunicativo, ruídos, barreiras, linguagem e feedback.", "1,0"),
    ("Comunicação interpessoal", "Demonstra escuta ativa, empatia, assertividade, atendimento e participação em reunião.", "1,0"),
    ("Documentos profissionais", "E-mail, comunicado, memorando, ata e relato seguem finalidade, estrutura e registro adequado.", "2,0"),
    ("Qualidade da escrita", "Clareza, objetividade, coesão, coerência, paragrafação, pontuação, concordância e ortografia.", "1,0"),
    ("Comunicação digital", "Aplica ferramentas colaborativas, netiqueta, gestão de tarefas, segurança e privacidade.", "1,0"),
    ("Apresentação oral", "Fala clara, organizada, adequada ao contexto, com postura, tom de voz e resposta assertiva.", "1,5"),
    ("Autoria e organização", "Material autoral, completo, organizado, revisado e com fontes identificadas.", "0,5"),
)
for idx, (criterion, evidence, value) in enumerate(criteria):
    cells = rubric.add_row().cells
    for col, text in enumerate((criterion, evidence, value, "")):
        if idx % 2:
            set_fill(cells[col], "F4F7F8")
        p = cells[col].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col >= 2 else WD_ALIGN_PARAGRAPH.LEFT
        style_run(p.add_run(text), 8.7, bold=(col == 0), color=PETROLEO if col == 0 else "000000")
table_geometry(rubric, widths)

body(doc, "Nota final: ______ / 10,0", "Nota final")
body(doc, "Resultado:  ☐ Aprovado(a)  ☐ Necessita nova orientação", "Resultado")
body(doc, "Feedback do professor:")
for _ in range(4):
    p = doc.add_paragraph("_" * 92)
    p.paragraph_format.space_after = Pt(5)
    style_run(p.runs[0], 9, color="9AA5AA")

doc.core_properties.title = "Recuperação — Comunicação Oral e Escrita"
doc.core_properties.subject = "Portfólio individual abrangendo a ementa completa da UC"
doc.core_properties.author = "SENAI"
OUTPUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUTPUT)
print(OUTPUT)
