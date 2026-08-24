from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "INTRODUCAO_COMUNICACAO_ORAL_ESCRITA" / "AVALIACOES" / "AVALIACAO_PRATICA_COMUNICACAO_ORAL_ESCRITA.docx"
PETROLEO = "004D5C"
LARANJA = "FF6B35"
AZUL_CLARO = "EAF4F6"
CINZA = "5F6B73"


def set_cell_fill(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), color)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def style_run(run, size=10.5, bold=False, color="000000", italic=False):
    run.font.name = "Arial"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Arial")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    style_run(run, 15 if level == 1 else 12, True, PETROLEO)
    return p


def add_body(doc, text="", bold_label=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    if bold_label and text.startswith(bold_label):
        first, rest = text.split(":", 1)
        style_run(p.add_run(first + ":"), bold=True, color=PETROLEO)
        style_run(p.add_run(rest))
    else:
        style_run(p.add_run(text))
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(3)
    style_run(p.add_run(text))


def add_answer_lines(doc, count):
    for _ in range(count):
        p = doc.add_paragraph("_" * 92)
        p.paragraph_format.space_after = Pt(5)
        for run in p.runs:
            style_run(run, 9, color="9AA5AA")


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.65)
section.bottom_margin = Inches(0.65)
section.left_margin = Inches(0.75)
section.right_margin = Inches(0.75)
section.header_distance = Inches(0.3)
section.footer_distance = Inches(0.3)

normal = doc.styles["Normal"]
normal.font.name = "Arial"
normal.font.size = Pt(10.5)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.15
for style_name, size in (("Heading 1", 15), ("Heading 2", 12)):
    style = doc.styles[style_name]
    style.font.name = "Arial"
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor.from_string(PETROLEO)

header = section.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.CENTER
style_run(header.add_run("RIO DO SUL MAIS TECH · SENAI · Prefeitura Municipal de Rio do Sul"), 8.5, True, PETROLEO)

footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
style_run(footer.add_run("Introdução à Comunicação Oral e Escrita para o Mundo do Trabalho"), 8, color=CINZA)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(4)
p.paragraph_format.space_after = Pt(2)
style_run(p.add_run("AVALIAÇÃO PRÁTICA"), 22, True, PETROLEO)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(12)
style_run(p.add_run("Comunicação profissional: escrever, ouvir e apresentar"), 13, True, LARANJA)

info = doc.add_table(rows=2, cols=2)
info.style = "Table Grid"
set_table_geometry(info, [4680, 4680])
fields = (("Aluno(a):", "Turma:"), ("Data:", "Equipe/dupla:"))
for row, labels in zip(info.rows, fields):
    for cell, label in zip(row.cells, labels):
        set_cell_fill(cell, AZUL_CLARO)
        p = cell.paragraphs[0]
        style_run(p.add_run(f"{label} __________________________________"), 10, True, PETROLEO)

add_heading(doc, "1. Situação-problema", 1)
add_body(doc, "Você faz parte da equipe de atendimento da empresa fictícia EcoTech Jovem, que organiza oficinas de tecnologia para estudantes. Uma escola informou que recebeu uma mensagem confusa sobre a mudança de horário de uma oficina. Sua equipe deverá corrigir a comunicação e apresentar uma solução profissional.")

callout = doc.add_table(rows=1, cols=1)
callout.style = "Table Grid"
set_table_geometry(callout, [9360])
set_cell_fill(callout.cell(0, 0), "FFF2EB")
p = callout.cell(0, 0).paragraphs[0]
style_run(p.add_run("Mensagem recebida: "), 10, True, LARANJA)
style_run(p.add_run('“Oi, mudou o negócio de amanhã. Acho que vai ser mais tarde. Avisa o pessoal aí porque não sei direito. Valeu.”'), 10, italic=True)

add_heading(doc, "2. Entregas da avaliação", 1)
add_body(doc, "A avaliação será realizada em dupla e terá duas partes complementares. Tempo sugerido: 50 minutos para preparação e 3 a 4 minutos para apresentação.")
add_bullet(doc, "Parte A — Produção escrita: redigir um e-mail corporativo claro e completo.")
add_bullet(doc, "Parte B — Comunicação oral: apresentar a solução e simular o atendimento à escola.")
add_bullet(doc, "Valor total: 10,0 pontos. Nota mínima para aprovação: 7,0.")

add_heading(doc, "PARTE A — E-mail corporativo (5,0 pontos)", 1)
add_body(doc, "Transforme a mensagem recebida em um e-mail profissional. Considere que a oficina foi remarcada de terça-feira, às 14h, para quarta-feira, às 15h, no mesmo local. A escola deve confirmar o recebimento até as 12h do dia anterior.")
add_body(doc, "Seu e-mail deve conter:")
for item in ("assunto objetivo;", "saudação adequada;", "explicação clara da alteração;", "orientação para confirmação;", "despedida e assinatura da equipe."):
    add_bullet(doc, item)
add_body(doc, "Assunto: ______________________________________________________________________", "Assunto")
add_answer_lines(doc, 9)

doc.add_page_break()
add_heading(doc, "PARTE B — Apresentação e simulação oral (5,0 pontos)", 1)
add_body(doc, "Prepare uma apresentação de 3 a 4 minutos. Um integrante representará a EcoTech Jovem e o outro representará a escola. Depois, troquem brevemente os papéis para que ambos demonstrem fala e escuta.")
add_body(doc, "Roteiro obrigatório:")
for item in (
    "apresentar o problema sem culpar pessoas;",
    "explicar como o ruído prejudicou a comunicação;",
    "comunicar corretamente o novo dia e horário;",
    "usar escuta ativa: não interromper, confirmar o entendimento e responder com respeito;",
    "encerrar com uma solução e um pedido de confirmação."
):
    add_bullet(doc, item)

add_heading(doc, "Planejamento da dupla", 2)
prompts = (
    "Qual foi o principal ruído na mensagem original?",
    "Como a equipe explicará a mudança com clareza?",
    "Qual frase assertiva será utilizada?",
    "Como vocês confirmarão que a escola compreendeu?",
)
for prompt in prompts:
    add_body(doc, prompt)
    add_answer_lines(doc, 2)

add_heading(doc, "3. Autoavaliação do estudante", 1)
add_body(doc, "Marque uma opção em cada frase:")
for text in (
    "Consegui transformar uma mensagem confusa em comunicação profissional.",
    "Falei com clareza, respeito e postura adequada.",
    "Pratiquei escuta ativa durante a simulação.",
):
    add_bullet(doc, text)
    for option in (
        "☐ Sim",
        "☐ Parcialmente — Justifique: ________________________________________________",
        "☐ Não — Justifique: _______________________________________________________",
    ):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.75)
        p.paragraph_format.space_after = Pt(2)
        style_run(p.add_run(option), 10)

add_heading(doc, "4. Critérios de avaliação", 1)
rubric = doc.add_table(rows=1, cols=4)
rubric.style = "Table Grid"
set_table_geometry(rubric, [2800, 4120, 1220, 1220])
headers = ("Critério", "Evidência esperada", "Valor", "Nota")
for cell, label in zip(rubric.rows[0].cells, headers):
    set_cell_fill(cell, PETROLEO)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_run(p.add_run(label), 9.5, True, "FFFFFF")
rows = (
    ("Estrutura do e-mail", "Assunto, saudação, corpo, despedida e assinatura.", "1,0", ""),
    ("Clareza e objetividade", "Informações completas, organizadas e sem ambiguidades.", "1,5", ""),
    ("Linguagem escrita", "Registro formal, coesão, pontuação e ortografia adequadas.", "1,5", ""),
    ("Netiqueta e responsabilidade", "Tom respeitoso e orientação segura para confirmação.", "1,0", ""),
    ("Clareza da fala", "Voz audível, sequência lógica e informações corretas.", "1,5", ""),
    ("Postura e comunicação não verbal", "Contato visual, postura profissional e gestos adequados.", "1,0", ""),
    ("Assertividade", "Expõe o problema e a solução com respeito, sem agressividade.", "1,0", ""),
    ("Escuta ativa e trabalho em dupla", "Ouve, confirma entendimento, coopera e respeita turnos.", "1,5", ""),
)
for idx, values in enumerate(rows):
    cells = rubric.add_row().cells
    for col, (cell, value) in enumerate(zip(cells, values)):
        if idx % 2:
            set_cell_fill(cell, "F4F7F8")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col >= 2 else WD_ALIGN_PARAGRAPH.LEFT
        style_run(p.add_run(value), 8.8, bold=(col == 0), color=PETROLEO if col == 0 else "000000")
set_table_geometry(rubric, [2800, 4120, 1220, 1220])

add_body(doc, "Nota final: ______ / 10,0")
add_body(doc, "Feedback do professor:")
add_answer_lines(doc, 3)

OUTPUT.parent.mkdir(parents=True, exist_ok=True)
doc.core_properties.title = "Avaliação Prática — Comunicação Oral e Escrita"
doc.core_properties.subject = "Rio do Sul Mais Tech — SENAI"
doc.core_properties.author = "SENAI"
doc.save(OUTPUT)
print(OUTPUT)
